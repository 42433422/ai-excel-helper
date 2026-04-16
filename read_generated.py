from docx import Document

# 读取生成的文件
doc = Document(r"e:\FHD\424\outputs\sales_contracts\销售合同_深圳市百木鼎家具有限公司_20260412_191728.docx")

print("=== 段落内容 ===")
for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        print(f"[{i}] {para.text[:100]}")

print("\n=== 表格内容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx} ---")
    for r_idx, row in enumerate(table.rows[:5]):
        cells = [cell.text[:30] for cell in row.cells]
        print(f"行 {r_idx}: {cells}")
