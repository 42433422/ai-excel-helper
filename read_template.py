from docx import Document

# 读取模板文件
doc = Document(r"e:\FHD\424\PZMOB.docx")

print("=== 段落内容 ===")
for i, para in enumerate(doc.paragraphs[:30]):
    if para.text.strip():
        print(f"[{i}] {para.text[:100]}")

print(f"\n总段落数: {len(doc.paragraphs)}")
print(f"总表格数: {len(doc.tables)}")

print("\n=== 表格内容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx} (行数: {len(table.rows)}) ---")
    for r_idx, row in enumerate(table.rows[:5]):
        cells = [cell.text[:30] for cell in row.cells]
        print(f"行 {r_idx}: {cells}")
