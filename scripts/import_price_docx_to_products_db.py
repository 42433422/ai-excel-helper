# -*- coding: utf-8 -*-
"""
从 424 目录下「转 Word_扫描全能王 … 13.22.docx」类报价表导入：
- 客户名称 → purchase_units（若不存在）
- 表格行 → products（model_number, name, specification, price, unit）

默认写入 FHD_PRODUCTS_DB 或 XCAGI/products.db。

用法:
  python scripts/import_price_docx_to_products_db.py
  python scripts/import_price_docx_to_products_db.py --dry-run
  python scripts/import_price_docx_to_products_db.py --file "E:\\FHD\\424\\xxx.docx"
"""

from __future__ import annotations

import argparse
import os
import re
import sqlite3
import sys
import zipfile
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]


def _default_docx() -> Path:
    d424 = REPO / "424"
    hits = sorted(d424.glob("*11.04.26*13.22*.docx"))
    if hits:
        return hits[0]
    hits = sorted(d424.glob("*.docx"))
    for h in hits:
        if "template" in h.name.lower():
            continue
        return h
    raise FileNotFoundError("未在 424 下找到 docx")


def _extract_customer_and_rows_from_docx(path: Path) -> tuple[str, list[dict]]:
    from docx import Document

    doc = Document(str(path))
    customer = ""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if "客户名称" in t or "客户名称：" in t:
            m = re.search(
                r"客户名称[：:]\s*([^\t\n报价日期]+)",
                t.replace("\u3000", " "),
            )
            if m:
                customer = m.group(1).strip()
            else:
                parts = re.split(r"[\t]+|报价日期", t)
                if len(parts) >= 2:
                    customer = (
                        parts[0]
                        .replace("客户名称：", "")
                        .replace("客户名称:", "")
                        .strip()
                    )
            break

    rows_out: list[dict] = []
    if not doc.tables:
        return customer, rows_out
    table = doc.tables[0]
    header = [c.text.strip() for c in table.rows[0].cells]
    # 期望: 编号 产品名称 规格 调价前 调价后 备注
    col_idx = {h: i for i, h in enumerate(header)}
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 5:
            continue
        code = cells[col_idx.get("编号", 0)] if "编号" in col_idx else cells[0]
        name = cells[col_idx.get("产品名称", 1)] if "产品名称" in col_idx else cells[1]
        spec = cells[col_idx.get("规格", 2)] if "规格" in col_idx else cells[2]
        price_after = cells[col_idx.get("调价后", 4)] if "调价后" in col_idx else cells[4]
        code = (code or "").strip()
        name = (name or "").strip()
        spec = (spec or "").strip()
        if not name and not code:
            continue
        price_val = _parse_price_yuan_per_kg(price_after)
        rows_out.append(
            {
                "model_number": code or name[:40],
                "name": name or code,
                "specification": spec,
                "price": price_val,
            }
        )
    return customer, rows_out


def _parse_price_yuan_per_kg(s: str) -> float:
    s = (s or "").strip().replace("／", "/")
    m = re.search(r"([\d.]+)\s*元", s)
    if not m:
        return 0.0
    try:
        return float(m.group(1))
    except ValueError:
        return 0.0


def _resolve_db_path() -> Path:
    env = os.environ.get("FHD_PRODUCTS_DB", "").strip()
    if env:
        p = Path(env).expanduser()
        if p.is_file():
            return p.resolve()
    for d in (REPO / "XCAGI", REPO / "xcagi"):
        cand = d / "products.db"
        if cand.is_file():
            return cand.resolve()
    raise FileNotFoundError("未找到 products.db，请设置 FHD_PRODUCTS_DB")


def _ensure_schema(conn: sqlite3.Connection) -> None:
    cur = conn.cursor()
    cur.execute(
        "SELECT 1 FROM sqlite_master WHERE type='table' AND name='purchase_units'"
    )
    if not cur.fetchone():
        cur.execute(
            """
            CREATE TABLE purchase_units (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                unit_name TEXT NOT NULL,
                contact_person TEXT,
                contact_phone TEXT,
                address TEXT,
                is_active INTEGER NOT NULL DEFAULT 1,
                created_at TEXT,
                updated_at TEXT
            )
            """
        )
    cur.execute("SELECT 1 FROM sqlite_master WHERE type='table' AND name='products'")
    if not cur.fetchone():
        raise RuntimeError("products 表不存在，拒绝创建以免结构不一致")


def import_data(db_path: Path, customer: str, rows: list[dict], dry_run: bool) -> None:
    if not customer:
        raise ValueError("未能从 Word 解析到客户名称（购买单位）")
    if not rows:
        raise ValueError("未能从 Word 表格解析到产品行")

    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    _ensure_schema(conn)
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(products)")
    pcols = {str(r[1]) for r in cur.fetchall()}
    required = {"model_number", "name"}
    if not required.issubset(pcols):
        raise RuntimeError(f"products 表缺少列: {required - pcols}")

    cur.execute(
        "SELECT id FROM purchase_units WHERE unit_name = ? AND (is_active IS NULL OR is_active = 1)",
        (customer,),
    )
    pu = cur.fetchone()
    if pu:
        print(f"购买单位已存在 id={pu['id']}: {customer}")
    else:
        print(f"新增购买单位: {customer}")
        if not dry_run:
            cur.execute(
                """
                INSERT INTO purchase_units (unit_name, contact_person, contact_phone, address, is_active)
                VALUES (?, '', '', '', 1)
                """,
                (customer,),
            )

    inserted = 0
    skipped = 0
    for r in rows:
        model = r["model_number"]
        name = r["name"]
        # 型号在全局未必唯一：同一型号可对应不同购买单位（unit），故按 (model_number, unit) 判重
        if "unit" in pcols:
            cur.execute(
                """
                SELECT id FROM products
                WHERE model_number = ? AND TRIM(unit) = ?
                  AND (is_active IS NULL OR is_active = 1)
                LIMIT 1
                """,
                (model, customer),
            )
        else:
            cur.execute(
                "SELECT id FROM products WHERE model_number = ? AND (is_active IS NULL OR is_active = 1) LIMIT 1",
                (model,),
            )
        if cur.fetchone():
            print(f"  跳过(该客户下型号已存在): {model} {name}")
            skipped += 1
            continue
        cols = ["model_number", "name"]
        vals: list = [model, name]
        if "specification" in pcols:
            cols.append("specification")
            vals.append(r.get("specification") or "")
        if "price" in pcols:
            cols.append("price")
            vals.append(r.get("price") or 0)
        if "unit" in pcols:
            cols.append("unit")
            vals.append(customer)
        if "quantity" in pcols:
            cols.append("quantity")
            vals.append(0)
        if "description" in pcols:
            cols.append("description")
            vals.append("")
        if "category" in pcols:
            cols.append("category")
            vals.append("")
        if "brand" in pcols:
            cols.append("brand")
            vals.append("")
        if "is_active" in pcols:
            cols.append("is_active")
            vals.append(1)
        ph = ",".join(["?"] * len(vals))
        sql = f"INSERT INTO products ({', '.join(cols)}) VALUES ({ph})"
        print(f"  新增产品: {model} | {name} | {r.get('specification','')}")
        if not dry_run:
            cur.execute(sql, vals)
        inserted += 1

    if not dry_run:
        conn.commit()
    conn.close()
    print(f"完成 dry_run={dry_run} 新增产品行={inserted} 跳过={skipped}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--file", type=Path, default=None, help="docx 路径")
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()
    docx = args.file or _default_docx()
    if not docx.is_file():
        print("文件不存在:", docx, file=sys.stderr)
        return 1
    customer, rows = _extract_customer_and_rows_from_docx(docx)
    db_path = _resolve_db_path()
    print("DOCX:", docx)
    print("DB:", db_path)
    print("客户(购买单位):", customer)
    print("表格行数:", len(rows))
    import_data(db_path, customer, rows, args.dry_run)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
