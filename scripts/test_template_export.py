"""
测试模板导出功能

验证模板上传、解析、填充和导出功能是否正常工作。
"""

import sys
import tempfile
from pathlib import Path
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from backend.excel_template import handle_excel_template
from backend.word_template import handle_word_template
from backend.template_upload import save_uploaded_file, get_upload_dir
from backend.template_database import init_db, get_template_by_id, list_templates


def create_test_excel_template() -> Path:
    """创建测试用 Excel 模板"""
    wb = Workbook()
    ws = wb.active
    ws.title = "报价单"
    
    # 添加表头
    headers = ["产品名称", "型号", "数量", "单价", "总价"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = cell.font.copy(bold=True)
    
    # 添加示例数据行（作为模板）
    ws.cell(row=2, column=1, value="{{product_name}}")
    ws.cell(row=2, column=2, value="{{model_number}}")
    ws.cell(row=2, column=3, value="{{quantity}}")
    ws.cell(row=2, column=4, value="{{unit_price}}")
    ws.cell(row=2, column=5, value="{{total_price}}")
    
    # 添加客户信息占位符
    ws.cell(row=10, column=1, value="客户名称：{{customer_name}}")
    ws.cell(row=11, column=1, value="报价日期：{{quote_date}}")
    ws.cell(row=12, column=1, value="报价单号：{{quote_number}}")
    
    # 保存
    upload_dir = get_upload_dir() / "excel"
    upload_dir.mkdir(parents=True, exist_ok=True)
    template_path = upload_dir / "test_template.xlsx"
    wb.save(template_path)
    wb.close()
    
    print(f"✓ 创建 Excel 测试模板：{template_path}")
    return template_path


def create_test_word_template() -> Path:
    """创建测试用 Word 模板"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("报价单", 0)
    
    # 添加客户信息段落
    doc.add_paragraph("客户名称：{{customer_name}}")
    doc.add_paragraph("报价日期：{{quote_date}}")
    doc.add_paragraph("报价单号：{{quote_number}}")
    
    # 添加表格
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ["产品名称", "型号", "数量", "单价", "总价"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    
    # 数据行（使用占位符）
    data_row = ["{{product_name}}", "{{model_number}}", "{{quantity}}", "{{unit_price}}", "{{total_price}}"]
    for i, value in enumerate(data_row):
        table.cell(1, i).text = value
    
    # 保存
    upload_dir = get_upload_dir() / "word"
    upload_dir.mkdir(parents=True, exist_ok=True)
    template_path = upload_dir / "test_template.docx"
    doc.save(template_path)
    
    print(f"✓ 创建 Word 测试模板：{template_path}")
    return template_path


def test_excel_parse(template_path: Path) -> bool:
    """测试 Excel 模板解析"""
    print("\n=== 测试 Excel 模板解析 ===")
    
    arguments = {
        "action": "parse",
        "file_path": str(template_path),
    }
    
    result = handle_excel_template(arguments)
    
    if "error" in result:
        print(f"✗ 解析失败：{result['error']}")
        return False
    
    print(f"✓ 解析成功")
    print(f"  工作表数量：{len(result.get('sheets', []))}")
    print(f"  占位符数量：{len(result.get('placeholders', []))}")
    
    if 'placeholders' in result:
        # 占位符可能是字典格式，需要提取名称
        placeholders = result['placeholders']
        if placeholders and isinstance(placeholders[0], dict):
            placeholder_names = [p.get('name', str(p)) for p in placeholders]
        else:
            placeholder_names = [str(p) for p in placeholders]
        print(f"  占位符列表：{', '.join(placeholder_names[:5])}")
    
    return True


def test_excel_fill(template_path: Path) -> bool:
    """测试 Excel 模板填充"""
    print("\n=== 测试 Excel 模板填充 ===")
    
    test_data = {
        "customer_name": "测试客户",
        "quote_date": "2026-04-12",
        "quote_number": "QT-2026-001",
        "product_name": "测试产品 A",
        "model_number": "MODEL-001",
        "quantity": "100",
        "unit_price": "10.5",
        "total_price": "1050.0",
    }
    
    arguments = {
        "action": "fill",
        "file_path": str(template_path),
        "data": test_data,
    }
    
    result = handle_excel_template(arguments)
    
    if "error" in result:
        print(f"✗ 填充失败：{result['error']}")
        return False
    
    output_path = result.get("output_path")
    if output_path:
        print(f"✓ 填充成功")
        print(f"  输出文件：{output_path}")
        
        # 验证输出文件存在
        if Path(output_path).exists():
            print(f"  ✓ 输出文件已创建")
            return True
        else:
            print(f"  ✗ 输出文件不存在")
            return False
    else:
        print(f"✗ 未返回输出路径")
        return False


def test_word_parse(template_path: Path) -> bool:
    """测试 Word 模板解析"""
    print("\n=== 测试 Word 模板解析 ===")
    
    arguments = {
        "action": "parse",
        "file_path": str(template_path),
    }
    
    result = handle_word_template(arguments)
    
    if "error" in result:
        print(f"✗ 解析失败：{result['error']}")
        return False
    
    print(f"✓ 解析成功")
    print(f"  表格数量：{result.get('structure', {}).get('total_tables', 0)}")
    print(f"  占位符数量：{result.get('structure', {}).get('total_placeholders', 0)}")
    
    placeholders = result.get('structure', {}).get('placeholders', [])
    if placeholders:
        print(f"  占位符列表：{', '.join(placeholders[:5])}")
    
    return True


def test_word_fill(template_path: Path) -> bool:
    """测试 Word 模板填充"""
    print("\n=== 测试 Word 模板填充 ===")
    
    test_data = {
        "customer_name": "测试客户",
        "quote_date": "2026-04-12",
        "quote_number": "QT-2026-001",
        "product_name": "测试产品 A",
        "model_number": "MODEL-001",
        "quantity": "100",
        "unit_price": "10.5",
        "total_price": "1050.0",
    }
    
    arguments = {
        "action": "fill",
        "file_path": str(template_path),
        "data": test_data,
    }
    
    result = handle_word_template(arguments)
    
    if "error" in result:
        print(f"✗ 填充失败：{result['error']}")
        return False
    
    output_path = result.get("output_path")
    if output_path:
        print(f"✓ 填充成功")
        print(f"  输出文件：{output_path}")
        
        # 验证输出文件存在
        if Path(output_path).exists():
            print(f"  ✓ 输出文件已创建")
            return True
        else:
            print(f"  ✗ 输出文件不存在")
            return False
    else:
        print(f"✗ 未返回输出路径")
        return False


def test_export_to_excel():
    """测试数据导出为 Excel"""
    print("\n=== 测试数据导出为 Excel ===")
    
    test_records = [
        {"产品名称": "产品 A", "型号": "MODEL-001", "数量": 100, "单价": 10.5, "总价": 1050.0},
        {"产品名称": "产品 B", "型号": "MODEL-002", "数量": 200, "单价": 20.5, "总价": 4100.0},
        {"产品名称": "产品 C", "型号": "MODEL-003", "数量": 150, "单价": 15.5, "总价": 2325.0},
    ]
    
    # 创建一个简单的模板文件用于导出
    upload_dir = get_upload_dir() / "excel"
    upload_dir.mkdir(parents=True, exist_ok=True)
    template_path = upload_dir / "export_template.xlsx"
    
    wb = Workbook()
    ws = wb.active
    ws.title = "导出数据"
    headers = list(test_records[0].keys())
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
    wb.save(template_path)
    wb.close()
    
    arguments = {
        "action": "export",
        "file_path": str(template_path),
        "records": test_records,
        "columns": ["产品名称", "型号", "数量", "单价", "总价"],
    }
    
    result = handle_excel_template(arguments)
    
    if "error" in result:
        print(f"✗ 导出失败：{result['error']}")
        return False
    
    output_path = result.get("output_path")
    if output_path:
        print(f"✓ 导出成功")
        print(f"  输出文件：{output_path}")
        print(f"  记录数：{len(test_records)}")
        
        # 验证输出文件存在
        if Path(output_path).exists():
            print(f"  ✓ 输出文件已创建")
            return True
        else:
            print(f"  ✗ 输出文件不存在")
            return False
    else:
        print(f"✗ 未返回输出路径")
        return False


def test_template_database():
    """测试模板数据库功能"""
    print("\n=== 测试模板数据库 ===")
    
    try:
        init_db()
        print(f"✓ 数据库初始化成功")
        
        templates = list_templates(limit=10)
        print(f"✓ 查询模板列表成功")
        print(f"  模板数量：{len(templates)}")
        
        if templates:
            first_template = templates[0]
            template = get_template_by_id(first_template.id)
            if template:
                print(f"✓ 查询模板详情成功")
                print(f"  模板名称：{template.name}")
                print(f"  模板类型：{template.type}")
            else:
                print(f"✗ 查询模板详情失败")
        
        return True
    except Exception as e:
        print(f"✗ 数据库测试失败：{e}")
        return False


def main():
    """主测试函数"""
    print("=" * 70)
    print("模板导出功能测试")
    print("=" * 70)
    
    # 创建测试模板
    excel_template = create_test_excel_template()
    word_template = create_test_word_template()
    
    # 运行测试
    results = {
        "Excel 解析": test_excel_parse(excel_template),
        "Excel 填充": test_excel_fill(excel_template),
        "Word 解析": test_word_parse(word_template),
        "Word 填充": test_word_fill(word_template),
        "Excel 导出": test_export_to_excel(),
        "数据库": test_template_database(),
    }
    
    # 打印结果摘要
    print("\n" + "=" * 70)
    print("测试结果摘要")
    print("=" * 70)
    
    passed = sum(1 for v in results.values() if v)
    total = len(results)
    
    for test_name, result in results.items():
        status = "✓ 通过" if result else "✗ 失败"
        print(f"{test_name:15s} {status}")
    
    print(f"\n总计：{passed}/{total} 测试通过")
    
    if passed == total:
        print("\n🎉 所有测试通过！模板导出功能正常工作。")
        return 0
    else:
        print(f"\n⚠ {total - passed} 个测试失败，请检查相关功能。")
        return 1


if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)
