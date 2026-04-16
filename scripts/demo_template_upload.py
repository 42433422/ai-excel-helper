"""
演示脚本：展示模板上传功能的使用

运行此脚本前，请确保：
1. 后端服务正在运行 (uvicorn backend.http_app:app --reload)
2. 已安装 requests 库 (pip install requests)
"""

import requests
import tempfile
from pathlib import Path
from openpyxl import Workbook
from docx import Document

BASE_URL = "http://localhost:8000"


def create_sample_excel() -> Path:
    """创建示例 Excel 模板"""
    wb = Workbook()
    ws = wb.active
    ws.title = "报价单"
    
    # 添加表头
    headers = ["产品名称", "型号", "数量", "单价", "总价", "备注"]
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
    
    # 添加示例数据
    ws.cell(row=2, column=1, value="产品 A")
    ws.cell(row=2, column=2, value="MODEL-001")
    ws.cell(row=2, column=3, value=100)
    ws.cell(row=2, column=4, value=10.5)
    ws.cell(row=2, column=5, value=1050.0)
    ws.cell(row=2, column=6, value="标准产品")
    
    # 添加占位符
    ws.cell(row=10, column=1, value="客户名称：{{customer_name}}")
    ws.cell(row=11, column=1, value="报价日期：{{quote_date}}")
    ws.cell(row=12, column=1, value="报价单号：{{quote_number}}")
    
    # 保存到临时文件
    tmp_path = Path(tempfile.mktemp(suffix=".xlsx"))
    wb.save(tmp_path)
    wb.close()
    
    return tmp_path


def create_sample_word() -> Path:
    """创建示例 Word 模板"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("报价单", 0)
    
    # 添加段落和占位符
    doc.add_paragraph("客户名称：{{customer_name}}")
    doc.add_paragraph("日期：{{quote_date}}")
    doc.add_paragraph("编号：{{quote_number}}")
    
    # 添加表格
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ["产品名称", "型号", "数量", "单价", "总价"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    # 示例数据
    table.cell(1, 0).text = "产品 A"
    table.cell(1, 1).text = "MODEL-001"
    table.cell(1, 2).text = "100"
    table.cell(1, 3).text = "10.5"
    table.cell(1, 4).text = "1050"
    
    # 保存到临时文件
    tmp_path = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(tmp_path)
    
    return tmp_path


def demo_upload_excel():
    """演示上传 Excel 模板"""
    print("\n=== 演示 1: 上传 Excel 模板 ===")
    
    excel_path = create_sample_excel()
    print(f"创建示例 Excel 文件：{excel_path}")
    
    try:
        with open(excel_path, 'rb') as f:
            files = {'file': f}
            data = {
                'type': 'excel',
                'name': '示例报价单模板',
                'description': '演示用的 Excel 报价单模板'
            }
            
            response = requests.post(
                f"{BASE_URL}/api/templates/upload",
                files=files,
                data=data
            )
        
        if response.status_code == 201:
            result = response.json()
            print(f"✓ 上传成功!")
            print(f"  模板 ID: {result['template_id']}")
            print(f"  模板类型：{result['type']}")
            print(f"  状态：{result['status']}")
            print(f"  识别的字段数：{len(result['fields'])}")
            
            if result['fields']:
                print(f"\n  前 3 个字段:")
                for i, field in enumerate(result['fields'][:3]):
                    print(f"    {i+1}. {field['display_name']} ({field['field_type']})")
            
            if 'metadata' in result and 'placeholders' in result['metadata']:
                placeholders = result['metadata']['placeholders']
                if placeholders:
                    print(f"\n  识别的占位符：{', '.join(placeholders)}")
            
            return result['template_id']
        else:
            print(f"✗ 上传失败：{response.status_code}")
            print(f"  错误信息：{response.json()}")
            return None
    
    finally:
        # 清理临时文件
        if excel_path.exists():
            excel_path.unlink()


def demo_upload_word():
    """演示上传 Word 模板"""
    print("\n=== 演示 2: 上传 Word 模板 ===")
    
    word_path = create_sample_word()
    print(f"创建示例 Word 文件：{word_path}")
    
    try:
        with open(word_path, 'rb') as f:
            files = {'file': f}
            data = {
                'type': 'word',
                'name': '示例合同模板',
                'description': '演示用的 Word 合同模板'
            }
            
            response = requests.post(
                f"{BASE_URL}/api/templates/upload",
                files=files,
                data=data
            )
        
        if response.status_code == 201:
            result = response.json()
            print(f"✓ 上传成功!")
            print(f"  模板 ID: {result['template_id']}")
            print(f"  模板类型：{result['type']}")
            print(f"  状态：{result['status']}")
            print(f"  识别的字段数：{len(result['fields'])}")
            
            if result['fields']:
                print(f"\n  字段列表:")
                for i, field in enumerate(result['fields'][:5]):
                    print(f"    {i+1}. {field['display_name']} ({field['field_type']})")
            
            return result['template_id']
        else:
            print(f"✗ 上传失败：{response.status_code}")
            print(f"  错误信息：{response.json()}")
            return None
    
    finally:
        # 清理临时文件
        if word_path.exists():
            word_path.unlink()


def demo_list_templates():
    """演示获取模板列表"""
    print("\n=== 演示 3: 获取模板列表 ===")
    
    response = requests.get(f"{BASE_URL}/api/templates?limit=5")
    
    if response.status_code == 200:
        result = response.json()
        templates = result['templates']
        
        print(f"✓ 获取成功!")
        print(f"  模板总数：{result['total']}")
        print(f"  返回数量：{len(templates)}")
        
        if templates:
            print(f"\n  模板列表:")
            for i, t in enumerate(templates[:3], 1):
                print(f"    {i}. {t['name']} ({t['type']}) - {t['status']}")
    else:
        print(f"✗ 获取失败：{response.status_code}")


def demo_get_template(template_id: str):
    """演示获取模板详情"""
    print(f"\n=== 演示 4: 获取模板详情 (ID: {template_id}) ===")
    
    response = requests.get(f"{BASE_URL}/api/templates/{template_id}")
    
    if response.status_code == 200:
        template = response.json()
        print(f"✓ 获取成功!")
        print(f"  名称：{template['name']}")
        print(f"  类型：{template['type']}")
        print(f"  状态：{template['status']}")
        print(f"  描述：{template.get('description', '无')}")
        print(f"  创建时间：{template['created_at']}")
        print(f"  字段数：{len(template['fields'])}")
    else:
        print(f"✗ 获取失败：{response.status_code}")


def demo_update_template(template_id: str):
    """演示更新模板信息"""
    print(f"\n=== 演示 5: 更新模板信息 ===")
    
    update_data = {
        'name': '更新后的模板名称',
        'description': '这是更新后的描述信息',
        'status': 'active'
    }
    
    response = requests.put(
        f"{BASE_URL}/api/templates/{template_id}",
        json=update_data
    )
    
    if response.status_code == 200:
        result = response.json()
        print(f"✓ 更新成功!")
        print(f"  新名称：{result['name']}")
        print(f"  新描述：{result['description']}")
        print(f"  新状态：{result['status']}")
    else:
        print(f"✗ 更新失败：{response.status_code}")
        print(f"  错误信息：{response.json()}")


def main():
    """主函数"""
    print("=" * 60)
    print("模板上传功能演示")
    print("=" * 60)
    
    # 检查服务是否可用
    try:
        response = requests.get(f"{BASE_URL}/docs")
        if response.status_code != 200:
            print(f"⚠ 警告：后端服务可能未运行 (状态码：{response.status_code})")
            print(f"  请先运行：uvicorn backend.http_app:app --reload")
            return
    except requests.exceptions.ConnectionError:
        print(f"✗ 错误：无法连接到后端服务 ({BASE_URL})")
        print(f"  请确保后端服务正在运行：uvicorn backend.http_app:app --reload")
        return
    
    print("✓ 后端服务连接成功")
    
    # 执行演示
    excel_template_id = demo_upload_excel()
    word_template_id = demo_upload_word()
    
    demo_list_templates()
    
    if excel_template_id:
        demo_get_template(excel_template_id)
        demo_update_template(excel_template_id)
    
    print("\n" + "=" * 60)
    print("演示完成!")
    print("=" * 60)
    print("\n提示:")
    print("- 查看 API 文档：http://localhost:8000/docs")
    print("- 查看模板列表：curl http://localhost:8000/api/templates")
    print("- 详细使用指南：docs/TEMPLATE_UPLOAD_GUIDE.md")


if __name__ == "__main__":
    main()
