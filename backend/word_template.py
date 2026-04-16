"""
Word (.docx) template handling for price table export.
"""

from __future__ import annotations

import json
import os
import re
import uuid
from datetime import date
from pathlib import Path
from typing import Any, Mapping

import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


_ALLOWED_SUFFIX = {".docx"}


def _load_units_from_db() -> dict[str, dict]:
    """从数据库加载采购单位信息"""
    from backend.product_db_read import load_purchase_units_for_templates

    return load_purchase_units_for_templates()


def _load_products_from_db() -> dict[str, dict]:
    """从数据库加载产品信息"""
    from backend.product_db_read import load_products_dict_for_templates

    return load_products_dict_for_templates()


def _try_semantic_match(user_input: str) -> dict | None:
    """尝试使用语义匹配找到产品"""
    try:
        from backend.product_semantic_matcher import semantic_match_product
        return semantic_match_product(user_input)
    except Exception:
        return None


def _validate_products_batch(products: list[dict[str, Any]]) -> dict[str, Any]:
    """批量校验产品列表"""
    try:
        from backend.contract_validator import validate_contract
        customer_name = None
        for p in products:
            if p.get("customer_name"):
                customer_name = p.get("customer_name")
                break
        if customer_name:
            return validate_contract(customer_name, products)
    except Exception:
        pass
    return {"valid": False, "message": "校验服务不可用"}


def _normalize_text(text: str) -> str:
    """文本标准化：去除多余空格和特殊字符"""
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def _fuzzy_match(text: str, candidates: dict[str, Any], threshold: float = 0.8) -> tuple[str, Any] | None:
    """模糊匹配文本到候选列表"""
    if not text:
        return None
    normalized = _normalize_text(text).lower()
    normalized_key = normalized.replace(' ', '')

    for key, value in candidates.items():
        key_normalized = key.lower().replace(' ', '')
        if normalized == key or normalized_key == key_normalized:
            return (key, value)

        if normalized in key.lower() or key.lower() in normalized:
            return (key, value)

    text_chars = set(normalized.replace(' ', ''))
    best_match = None
    best_score = 0

    for key, value in candidates.items():
        key_chars = set(key.lower().replace(' ', ''))
        if not key_chars:
            continue
        intersection = text_chars & key_chars
        score = len(intersection) / max(len(text_chars), len(key_chars))
        if score >= threshold and score > best_score:
            best_score = score
            best_match = (key, value)

    return best_match


def cleanup_price_data(
    raw_data: list[dict[str, Any]],
    customer_name: str | None = None,
    quote_date: str | None = None,
    skip_validation: bool = False,
) -> dict[str, Any]:
    """
    清理价格表数据：
    1. 标准化客户名称 -> 匹配数据库单位（优先语义匹配）
    2. 标准化产品名称 -> 匹配数据库产品（优先语义匹配，其次模糊匹配）
    3. 获取单价（从数据库或计算）
    4. 格式化日期为本日
    5. 校验产品型号是否存在，提供修正建议

    Args:
        raw_data: 原始产品数据列表
        customer_name: 客户名称
        quote_date: 报价日期
        skip_validation: 是否跳过校验（默认 False，进行完整校验）

    Returns:
        {
            "cleaned_data": [...],
            "unit_info": {...},
            "products_matched": {...},
            "unmatched_products": [...],
            "quote_date": "2026-04-11",
            "validation": {
                "valid": bool,
                "invalid_products": [...],
                "suggestions": {...}
            }
        }
    """
    units_db = _load_units_from_db()
    products_db = _load_products_from_db()

    today = date.today().strftime("%Y-%m-%d")
    final_date = quote_date or today

    matched_unit = None
    unmatched_unit = customer_name
    if customer_name:
        matched = _fuzzy_match(customer_name, units_db)
        if matched:
            matched_unit = matched[1]
            unmatched_unit = None

    cleaned_data = []
    unmatched_products = []
    products_matched = {}
    invalid_products = []
    suggestions = {}

    for row in raw_data:
        cleaned_row = dict(row)

        product_name_raw = str(row.get("产品名称") or row.get("产品") or row.get("name") or "")
        model_number_raw = str(row.get("型号") or row.get("model_number") or row.get("编号") or "")
        product_name = _normalize_text(product_name_raw)

        if not product_name and not model_number_raw:
            continue

        matched_info = None
        match_source = None

        if product_name:
            semantic_result = _try_semantic_match(product_name)
            if semantic_result:
                matched_info = semantic_result
                match_source = "semantic"
            else:
                product_match = _fuzzy_match(product_name, products_db)
                if product_match:
                    matched_info = product_match[1]
                    match_source = "fuzzy_name"

        if not matched_info and model_number_raw:
            semantic_result = _try_semantic_match(model_number_raw)
            if semantic_result:
                matched_info = semantic_result
                match_source = "semantic"
            else:
                for key, val in products_db.items():
                    if _normalize_text(key) == _normalize_text(model_number_raw) or _normalize_text(model_number_raw) in _normalize_text(key):
                        matched_info = val
                        match_source = "fuzzy_model"
                        break

        if matched_info:
            matched_key = matched_info.get("name") or matched_info.get("model_number", "")
            products_matched[product_name or model_number_raw] = {
                "matched_to": matched_key,
                "source": match_source,
                "score": matched_info.get("_match_score") if match_source == "semantic" else None,
            }
            cleaned_row["产品名称"] = matched_info.get("name", "")
            cleaned_row["型号"] = matched_info.get("model_number", "")
            cleaned_row["单位"] = matched_info.get("specification") or row.get("单位") or ""
            cleaned_row["单价"] = matched_info.get("price", 0)
            cleaned_row["产品ID"] = matched_info.get("id")
            cleaned_row["matched"] = True
            cleaned_row["_match_source"] = match_source
        else:
            unmatched_products.append(product_name or model_number_raw)
            cleaned_row["matched"] = False
            cleaned_row["_match_source"] = None

            if not skip_validation:
                invalid_products.append({
                    "input": product_name or model_number_raw,
                    "original": row,
                })

        cleaned_row["客户名称"] = matched_unit["unit_name"] if matched_unit else (customer_name or "")
        cleaned_row["报价日期"] = final_date

        cleaned_data.append(cleaned_row)

    result = {
        "cleaned_data": cleaned_data,
        "quote_date": final_date,
        "unit_info": matched_unit,
        "unmatched_unit": unmatched_unit,
        "products_matched": products_matched,
        "unmatched_products": list(set(unmatched_products)),
    }

    if matched_unit:
        result["unit_id"] = matched_unit.get("id")

    if not skip_validation and (invalid_products or unmatched_unit):
        validation_result = {
            "valid": len(invalid_products) == 0 and unmatched_unit is None,
            "invalid_products": invalid_products,
            "invalid_customer": unmatched_unit is not None,
            "message": "",
        }

        if unmatched_unit:
            validation_result["message"] = f"客户「{unmatched_unit}」不存在；"

        if invalid_products:
            validation_result["message"] += f"有 {len(invalid_products)} 个产品无法匹配：{[p['input'] for p in invalid_products[:3]]}"

        result["validation"] = validation_result
    else:
        result["validation"] = {"valid": True, "invalid_products": [], "message": "校验通过"}

    return result


def resolve_safe_word_path(workspace_root: str, file_path: str) -> Path:
    """
    Resolve file_path to an absolute path that must stay under workspace_root.
    """
    root = Path(workspace_root).resolve()
    if not root.is_dir():
        raise FileNotFoundError(f"workspace root is not a directory: {root}")

    raw = Path(file_path)
    if raw.is_absolute():
        candidate = raw.resolve()
    else:
        candidate = (root / raw).resolve()

    try:
        candidate.relative_to(root)
    except ValueError as e:
        raise PermissionError("file_path must resolve inside workspace root") from e

    if ".." in Path(file_path).parts:
        raise PermissionError("file_path must not contain '..'")

    if candidate.suffix.lower() not in _ALLOWED_SUFFIX:
        raise ValueError("only Word files (.docx) are supported")

    return candidate


def parse_word_table(path: Path) -> dict[str, Any]:
    """
    Parse all tables from a Word document.
    Returns a structure with table info for template mapping.
    """
    doc = Document(path)
    tables_data = []

    for idx, table in enumerate(doc.tables):
        rows_data = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        tables_data.append({
            "table_index": idx,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "rows": rows_data,
            "header": rows_data[0] if rows_data else [],
            "data_rows": rows_data[1:] if len(rows_data) > 1 else [],
        })

    return {
        "table_count": len(doc.tables),
        "tables": tables_data,
    }


def fill_word_template(
    template_path: Path,
    data: Mapping[str, Any],
    output_path: Path | None = None,
) -> dict[str, Any]:
    """
    Fill a Word template with data.
    - template_path: path to .docx template
    - data: dict with keys matching template placeholders or table data
    - output_path: if provided, save filled document here

    Template format: {{placeholder}} in paragraphs, or table cells matching data keys.
    If data contains a "products" key with a list of product dicts,
    the second table (products table) will be filled with product rows.
    """
    doc = Document(template_path)

    filled_count = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
                filled_count += 1
        if text != paragraph.text:
            paragraph.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell_text:
                        cell_text = cell_text.replace(placeholder, str(value))
                        filled_count += 1
                if cell_text != cell.text:
                    cell.text = cell_text

    products = data.get("products")
    if products and len(products) > 0 and len(doc.tables) >= 2:
        products_table = doc.tables[1]
        try:
            from copy import deepcopy
            from docx.oxml.ns import qn

            def _physical_row_count(table):
                return len(table._tbl.tr_lst)

            def _strip_vertical_merge_from_row(tr):
                for tc in tr.iter():
                    if tc.tag.endswith("}tcPr"):
                        for vm in tc.findall(".//" + qn("w:vMerge")):
                            vm.get(qn("w:val"))

            def _append_table_row_clone(table, template_row_index):
                n = _physical_row_count(table)
                if template_row_index < 0 or template_row_index >= n:
                    table.add_row()
                    return
                try:
                    src_tr = table.rows[template_row_index]._tr
                    new_tr = deepcopy(src_tr)
                    _strip_vertical_merge_from_row(new_tr)
                    table._tbl.append(new_tr)
                except Exception:
                    table.add_row()

            start_row = 1
            clone_src = start_row if start_row < _physical_row_count(products_table) else max(0, _physical_row_count(products_table) - 1)

            for i, product in enumerate(products):
                ri = start_row + i
                while _physical_row_count(products_table) <= ri:
                    _append_table_row_clone(products_table, clone_src)
                cells = products_table.rows[ri].cells
                if len(cells) >= 1:
                    cells[0].text = str(product.get("model_number") or "").strip()
                if len(cells) >= 2:
                    cells[1].text = str(product.get("name") or "").strip()
                if len(cells) >= 3:
                    cells[2].text = str(product.get("specification") or "").strip()
                if len(cells) >= 4:
                    cells[3].text = str(product.get("quantity", "")).strip()
                if len(cells) >= 5:
                    price = product.get("unit_price", "")
                    cells[4].text = str(price) if price else ""
                filled_count += 1

            for ri in range(start_row + len(products), len(products_table.rows)):
                cells = products_table.rows[ri].cells
                for j in range(min(5, len(cells))):
                    cells[j].text = ""
        except Exception as e:
            logger.warning(f"填充产品表格失败: {e}")

    if output_path:
        doc.save(output_path)

    return {
        "filled_count": filled_count,
        "output_path": str(output_path) if output_path else None,
    }


def export_table_to_word(
    df: pd.DataFrame,
    output_path: Path,
    include_header: bool = True,
) -> dict[str, Any]:
    """
    Export a DataFrame to a Word table.
    """
    doc = Document()
    table = doc.add_table(rows=len(df) + (1 if include_header else 0), cols=len(df.columns))

    if include_header:
        for col_idx, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[col_idx]
            cell.text = str(col_name)

    for row_idx, (_, row) in enumerate(df.iterrows()):
        target_row = row_idx + (1 if include_header else 0)
        for col_idx, value in enumerate(row):
            cell = table.rows[target_row].cells[col_idx]
            cell.text = str(value) if pd.notna(value) else ""

    doc.save(output_path)
    return {
        "row_count": len(df),
        "column_count": len(df.columns),
        "output_path": str(output_path),
    }


def _default_workspace_root() -> str:
    return os.environ.get("WORKSPACE_ROOT", os.getcwd())


def handle_word_template(
    arguments: Mapping[str, Any],
    workspace_root: str | None = None,
) -> dict[str, Any]:
    """
    Main entry point for Word template operations.
    Actions: parse, fill, export, cleanup
    """
    root = workspace_root or _default_workspace_root()
    action = (arguments.get("action") or "").lower()
    file_path = arguments.get("file_path")

    if action == "cleanup":
        raw_data = arguments.get("raw_data") or arguments.get("records") or []
        customer_name = arguments.get("customer_name")
        quote_date = arguments.get("quote_date")
        return cleanup_price_data(raw_data, customer_name=customer_name, quote_date=quote_date)

    if not file_path:
        return {"error": "missing_file_path"}

    try:
        path = resolve_safe_word_path(root, str(file_path))
    except (OSError, ValueError, PermissionError) as e:
        return {"error": "path_resolution_failed", "message": str(e)}

    if not path.is_file():
        return {"error": "file_not_found", "path": str(path)}

    if action == "parse":
        return parse_word_table(path)

    if action == "fill":
        data = arguments.get("data") or {}
        upload_dir = Path(root) / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        output_filename = f"filled_{uuid.uuid4().hex}.docx"
        output_path = upload_dir / output_filename

        result = fill_word_template(path, data, output_path)
        try:
            rel = output_path.relative_to(Path(root).resolve())
            result["output_path"] = rel.as_posix()
        except ValueError:
            result["output_path"] = str(output_path)

        return result

    if action == "export":
        records = arguments.get("records") or arguments.get("data_rows") or []
        columns = arguments.get("columns") or (list(records[0].keys()) if records else [])

        df = pd.DataFrame(records, columns=columns)

        upload_dir = Path(root) / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        output_filename = f"export_{uuid.uuid4().hex}.docx"
        output_path = upload_dir / output_filename

        result = export_table_to_word(df, output_path)
        try:
            rel = output_path.relative_to(Path(root).resolve())
            result["output_path"] = rel.as_posix()
        except ValueError:
            result["output_path"] = str(output_path)

        return result

    return {
        "error": "invalid_action",
        "action": action,
        "allowed": ["parse", "fill", "export"],
    }


def get_word_tool_registry() -> list[dict[str, Any]]:
    """
    Word template tool for OpenAI Chat Completions tools list.
    """
    return [
        {
            "type": "function",
            "function": {
                "name": "word_template",
                "description": (
                    "【Word 模板工具】处理 .docx 价格表模板："
                    "parse=解析表格结构；fill=用数据填充模板占位符；export=DataFrame导出为Word表格；"
                    "cleanup=清理价格表数据（自动校验客户名称和产品型号是否在数据库存在，"
                    "支持语义相似度匹配修正，提供校验结果和建议）"
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "action": {
                            "type": "string",
                            "enum": ["parse", "fill", "export", "cleanup"],
                            "description": "parse: 解析文档表格结构; fill: 填充模板占位符; export: DataFrame导出为Word; cleanup: 清理价格表数据",
                        },
                        "file_path": {"type": "string", "description": "Path to .docx file under workspace"},
                        "data": {
                            "type": "object",
                            "description": "For fill action: key-value pairs to replace {{placeholder}} in template",
                        },
                        "records": {
                            "type": "array",
                            "description": "For export/cleanup action: array of records",
                        },
                        "raw_data": {
                            "type": "array",
                            "description": "Alias for records, for cleanup action",
                        },
                        "customer_name": {
                            "type": "string",
                            "description": "For cleanup action: 客户名称 to match against database units",
                        },
                        "quote_date": {
                            "type": "string",
                            "description": "For cleanup action: 报价日期 (defaults to today)",
                        },
                        "columns": {
                            "type": "array",
                            "description": "For export action: column names for the Word table",
                        },
                        "data_rows": {
                            "type": "array",
                            "description": "Alias for records",
                        },
                    },
                    "required": ["action"],
                },
            },
        },
    ]