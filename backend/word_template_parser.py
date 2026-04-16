"""
Word template intelligent parser.
Automatically identifies document structure, paragraphs, tables, and placeholders.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_document_structure(doc: Document) -> dict[str, Any]:
    """
    Extract overall document structure.
    
    Returns:
        Document structure with sections, paragraphs, and tables.
    """
    result = {
        "sections": [],
        "tables": [],
        "placeholders": [],
        "styles_used": set(),
    }
    
    for section in doc.sections:
        section_info = {
            "name": section.name if hasattr(section, "name") else f"Section_{len(result['sections']) + 1}",
            "page_width": section.page_width.mm if section.page_width else None,
            "page_height": section.page_height.mm if section.page_height else None,
            "orientation": "landscape" if section.orientation == 1 else "portrait",
        }
        result["sections"].append(section_info)
    
    for element in doc.element.body:
        if element.tag.endswith("tbl"):
            table_obj = Table(element, doc)
            table_info = parse_word_table(table_obj, len(result["tables"]))
            result["tables"].append(table_info)
        elif element.tag.endswith("p"):
            para_obj = Paragraph(element, doc)
            parse_paragraph(para_obj, result)
    
    result["styles_used"] = list(result["styles_used"])
    result["placeholders"] = sorted(list(set(result["placeholders"])))
    
    return result


def parse_word_table(table: Table, index: int) -> dict[str, Any]:
    """
    Parse Word table structure.
    
    Returns:
        Table metadata with rows, columns, and headers.
    """
    table_info = {
        "index": index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "headers": [],
        "data": [],
        "has_header_row": False,
    }
    
    if len(table.rows) == 0:
        return table_info
    
    first_row = table.rows[0]
    header_cells = []
    
    for cell in first_row.cells:
        text = cell.text.strip()
        if text:
            header_cells.append(text)
    
    if len(header_cells) >= len(table.columns) * 0.5:
        table_info["has_header_row"] = True
        for i, text in enumerate(header_cells):
            table_info["headers"].append({
                "index": i,
                "display_name": text,
                "name": _normalize_field_name(text),
                "data_type": _infer_data_type_from_header(text),
            })
        
        for row_idx in range(1, len(table.rows)):
            row_data = []
            for col_idx, cell in enumerate(table.rows[row_idx].cells):
                text = cell.text.strip()
                if text:
                    row_data.append({
                        "column": col_idx,
                        "value": text,
                    })
            if row_data:
                table_info["data"].append(row_data)
    else:
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    row_data.append({
                        "column": col_idx,
                        "value": text,
                    })
            if row_data:
                table_info["data"].append(row_data)
    
    return table_info


def parse_paragraph(para: Paragraph, result: dict[str, Any]) -> None:
    """
    Parse paragraph and extract information.
    
    Updates result dict with placeholders and styles.
    """
    text = para.text.strip()
    
    if para.style and para.style.name:
        result["styles_used"].add(para.style.name)
    
    placeholders = extract_placeholders_from_text(text)
    if placeholders:
        result["placeholders"].extend(placeholders)
    
    if text:
        para_info = {
            "text": text,
            "style": para.style.name if para.style else None,
            "is_heading": para.style.name.startswith("Heading") if para.style else False,
            "heading_level": _extract_heading_level(para.style.name) if para.style and para.style.name.startswith("Heading") else None,
        }
        
        if "paragraphs" not in result:
            result["paragraphs"] = []
        result["paragraphs"].append(para_info)


def extract_placeholders_from_text(text: str) -> list[str]:
    """
    Extract placeholder patterns from text.
    
    Supports patterns:
    - {{field_name}}
    - {%field_name%}
    - ${field_name}
    """
    placeholders = []
    
    patterns = [
        r"\{\{([^}]+)\}\}",
        r"\{%([^%]+)%\}",
        r"\$\{([^}]+)\}",
        r"\[([^\]]+)\]",
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        placeholders.extend(matches)
    
    return placeholders


def _normalize_field_name(text: str) -> str:
    """Normalize field name for programmatic use."""
    text = text.strip().lower()
    text = re.sub(r"[\s\-\(\)]+", "_", text)
    text = re.sub(r"[^\w\u4e00-\u9fff]+", "", text)
    text = re.sub(r"_+", "_", text)
    text = text.strip("_")
    return text or "field"


def _infer_data_type_from_header(header: str) -> str:
    """Infer data type from header text."""
    header_lower = header.lower()
    
    if any(kw in header_lower for kw in ["日期", "date", "时间", "time"]):
        return "date"
    if any(kw in header_lower for kw in ["金额", "价格", "单价", "总价", "price", "amount", "cost"]):
        return "number"
    if any(kw in header_lower for kw in ["数量", "qty", "count", "num"]):
        return "integer"
    if any(kw in header_lower for kw in ["电话", "手机", "contact", "phone"]):
        return "phone"
    if "@" in header or "email" in header_lower:
        return "email"
    
    return "string"


def _extract_heading_level(style_name: str) -> int | None:
    """Extract heading level from style name."""
    match = re.search(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def extract_metadata(doc: Document) -> dict[str, Any]:
    """
    Extract document metadata.
    
    Returns:
        Document metadata including properties and statistics.
    """
    metadata = {
        "paragraph_count": 0,
        "table_count": 0,
        "character_count": 0,
        "word_count": 0,
        "core_properties": {},
    }
    
    try:
        if hasattr(doc, "core_properties"):
            props = doc.core_properties
            metadata["core_properties"] = {
                "title": props.title if props.title else None,
                "subject": props.subject if props.subject else None,
                "author": props.author if props.author else None,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "keywords": props.keywords if props.keywords else None,
            }
    except Exception:
        pass
    
    for para in doc.paragraphs:
        metadata["paragraph_count"] += 1
        text = para.text.strip()
        metadata["character_count"] += len(text)
        metadata["word_count"] += len(text.split())
    
    metadata["table_count"] = len(doc.tables)
    
    return metadata


def parse_word_template(file_path: Path) -> dict[str, Any]:
    """
    Main entry point for parsing Word template.
    
    Returns comprehensive template metadata.
    """
    doc = Document(file_path)
    
    structure = extract_document_structure(doc)
    metadata = extract_metadata(doc)
    
    result = {
        "file_name": file_path.name,
        "structure": structure,
        "metadata": metadata,
        "summary": {
            "total_sections": len(structure["sections"]),
            "total_tables": len(structure["tables"]),
            "total_paragraphs": len(structure.get("paragraphs", [])),
            "total_placeholders": len(structure["placeholders"]),
            "has_header_rows": sum(1 for t in structure["tables"] if t.get("has_header_row")),
        },
    }
    
    return result
