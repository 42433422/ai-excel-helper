"""
Tests for template upload and parsing functionality.
"""

import os
import tempfile
from pathlib import Path

import pytest
from openpyxl import Workbook
from docx import Document

from backend.template_upload import (
    validate_file_extension,
    validate_mime_type,
    generate_unique_filename,
    format_file_size,
)
from backend.excel_template_parser import (
    parse_excel_template,
    detect_table_regions,
    extract_headers,
    _normalize_field_name,
    _infer_data_type_from_header,
)
from backend.word_template_parser import (
    parse_word_template,
    extract_placeholders_from_text,
    _normalize_field_name as normalize_word_field_name,
)


class TestFileValidation:
    """Test file validation functions."""
    
    def test_validate_excel_extension(self):
        """Test Excel extension validation."""
        assert validate_file_extension("test.xlsx", "excel") is True
        assert validate_file_extension("test.xlsm", "excel") is True
        assert validate_file_extension("test.xls", "excel") is False
        assert validate_file_extension("test.docx", "excel") is False
    
    def test_validate_word_extension(self):
        """Test Word extension validation."""
        assert validate_file_extension("test.docx", "word") is True
        assert validate_file_extension("test.doc", "word") is False
        assert validate_file_extension("test.xlsx", "word") is False
    
    def test_validate_logo_extension(self):
        """Test logo extension validation."""
        assert validate_file_extension("logo.png", "logo") is True
        assert validate_file_extension("logo.jpg", "logo") is True
        assert validate_file_extension("logo.jpeg", "logo") is True
        assert validate_file_extension("logo.gif", "logo") is True
        assert validate_file_extension("logo.svg", "logo") is True
        assert validate_file_extension("logo.xlsx", "logo") is False
    
    def test_validate_mime_type(self):
        """Test MIME type validation."""
        assert validate_mime_type(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "excel",
        ) is True
        assert validate_mime_type(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "word",
        ) is True
        assert validate_mime_type("image/png", "logo") is True
        assert validate_mime_type("text/plain", "excel") is False
    
    def test_generate_unique_filename(self):
        """Test unique filename generation."""
        filename1 = generate_unique_filename("test.xlsx")
        filename2 = generate_unique_filename("test.xlsx")
        
        assert filename1.endswith(".xlsx")
        assert filename2.endswith(".xlsx")
        assert filename1 != filename2
        
        assert len(filename1) == 32 + 5
        assert filename1.replace(".xlsx", "").isalnum()
    
    def test_format_file_size(self):
        """Test file size formatting."""
        assert format_file_size(100) == "100.00 B"
        assert format_file_size(1024) == "1.00 KB"
        assert format_file_size(1048576) == "1.00 MB"
        assert format_file_size(1073741824) == "1.00 GB"


class TestExcelParser:
    """Test Excel template parser."""
    
    def test_normalize_field_name(self):
        """Test field name normalization."""
        assert _normalize_field_name("Product Name") == "product_name"
        assert _normalize_field_name("客户名称") == "客户名称"
        assert _normalize_field_name("Price (USD)") == "price_usd"
        assert _normalize_field_name("  数量  ") == "数量"
        assert _normalize_field_name("") == "field"
    
    def test_infer_data_type_from_header(self):
        """Test data type inference from headers."""
        assert _infer_data_type_from_header("日期") == "date"
        assert _infer_data_type_from_header("Date") == "date"
        assert _infer_data_type_from_header("价格") == "number"
        assert _infer_data_type_from_header("Price") == "number"
        assert _infer_data_type_from_header("数量") == "integer"
        assert _infer_data_type_from_header("Qty") == "integer"
        assert _infer_data_type_from_header("产品名称") == "string"
    
    def test_parse_simple_excel_template(self):
        """Test parsing a simple Excel template."""
        wb = Workbook()
        ws = wb.active
        ws.title = "报价单"
        
        headers = ["产品名称", "型号", "数量", "单价", "总价", "日期"]
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        ws.cell(row=2, column=1, value="产品 A")
        ws.cell(row=2, column=2, value="MODEL-001")
        ws.cell(row=2, column=3, value=100)
        ws.cell(row=2, column=4, value=10.5)
        ws.cell(row=2, column=5, value=1050.0)
        ws.cell(row=2, column=6, value="2024-01-15")
        
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            
            wb.save(tmp_path)
            wb.close()
            
            result = parse_excel_template(tmp_path)
            
            assert result["file_name"] == tmp_path.name
            assert len(result["sheets"]) == 1
            assert result["sheets"][0]["name"] == "报价单"
            assert len(result["sheets"][0]["headers"]) >= 5
            
            field_names = [h["name"] for h in result["sheets"][0]["headers"]]
            assert "产品名称" in str(field_names) or "product_name" in str(field_names)
            
        finally:
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except Exception:
                    pass
    
    def test_parse_excel_with_placeholders(self):
        """Test parsing Excel with placeholders."""
        wb = Workbook()
        ws = wb.active
        
        ws.cell(row=1, column=1, value="客户：{{customer_name}}")
        ws.cell(row=2, column=1, value="日期：{{quote_date}}")
        ws.cell(row=3, column=1, value="编号：{{quote_number}}")
        
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            
            wb.save(tmp_path)
            wb.close()
            
            result = parse_excel_template(tmp_path)
            
            assert "customer_name" in result["placeholders"]
            assert "quote_date" in result["placeholders"]
            assert "quote_number" in result["placeholders"]
            
        finally:
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except Exception:
                    pass


class TestWordParser:
    """Test Word template parser."""
    
    def test_extract_placeholders_double_braces(self):
        """Test placeholder extraction with {{}} pattern."""
        text = "尊敬的{{customer_name}}，您好！"
        placeholders = extract_placeholders_from_text(text)
        assert "customer_name" in placeholders
    
    def test_extract_placeholders_multiple(self):
        """Test multiple placeholder extraction."""
        text = "{{name}} 先生/女士，您的订单{{order_id}}已确认，金额{{amount}}元。"
        placeholders = extract_placeholders_from_text(text)
        assert len(placeholders) >= 3
        assert "name" in placeholders
        assert "order_id" in placeholders
        assert "amount" in placeholders
    
    def test_extract_placeholders_different_patterns(self):
        """Test placeholder extraction with different patterns."""
        text1 = "Hello {%name%}"
        text2 = "Value: ${value}"
        text3 = "Item: [item_name]"
        
        assert len(extract_placeholders_from_text(text1)) > 0
        assert len(extract_placeholders_from_text(text2)) > 0
        assert len(extract_placeholders_from_text(text3)) > 0
    
    def test_normalize_word_field_name(self):
        """Test Word field name normalization."""
        assert normalize_word_field_name("Field Name") == "field_name"
        assert normalize_word_field_name("字段名称") == "字段名称"
        assert normalize_word_field_name("Price(USD)") == "price_usd"
    
    def test_parse_simple_word_template(self):
        """Test parsing a simple Word template."""
        doc = Document()
        
        doc.add_heading("报价单", 0)
        doc.add_paragraph("客户名称：{{customer_name}}")
        doc.add_paragraph("日期：{{quote_date}}")
        
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "产品名称"
        table.cell(0, 1).text = "数量"
        table.cell(0, 2).text = "价格"
        table.cell(1, 0).text = "产品 A"
        table.cell(1, 1).text = "100"
        table.cell(1, 2).text = "1000"
        
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            
            doc.save(tmp_path)
            doc = None
            
            result = parse_word_template(tmp_path)
            
            assert result["file_name"] == tmp_path.name
            assert result["summary"]["total_placeholders"] >= 2
            assert result["summary"]["total_tables"] == 1
            assert "customer_name" in result["structure"]["placeholders"]
            assert "quote_date" in result["structure"]["placeholders"]
            
        finally:
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except Exception:
                    pass


class TestIntegration:
    """Integration tests for template system."""
    
    def test_template_upload_flow(self):
        """Test complete template upload flow."""
        from backend.template_database import init_db, get_db_url
        from sqlalchemy import create_engine, text
        
        db_url = get_db_url()
        engine = create_engine(db_url)
        
        init_db()
        
        with engine.connect() as conn:
            result = conn.execute(
                text("SELECT name FROM sqlite_master WHERE type='table' AND name='templates'")
            )
            row = result.fetchone()
            assert row is not None, "Templates table should exist"
    
    def test_database_initialization(self):
        """Test database initialization."""
        from backend.template_database import init_db, get_session, Template
        
        init_db()
        
        session = get_session()
        
        try:
            count = session.query(Template).count()
            assert count >= 0
        finally:
            session.close()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
