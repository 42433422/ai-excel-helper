import io
import json
import sqlite3

import pytest
from docx import Document
from starlette.testclient import TestClient

from backend import http_app
from backend.price_list_docx_export import format_specification_for_price_export

_skip_legacy_sqlite_file_tests = pytest.mark.skip(
    reason="业务库仅 PostgreSQL：不再支持临时 SQLite 文件 mock，待用 PG 集成用例替代"
)


@pytest.mark.parametrize(
    "spec,expected",
    [
        ("3kg", "3kg /卡"),
        ("4KG", "4KG /卡"),
        ("4.0公斤", "4.0公斤 /卡"),
        ("5kg", "5kg /缶"),
        ("10 千克", "10 千克 /缶"),
        ("12kg", "12kg /缶"),
        ("13.9kg", "13.9kg /缶"),
        ("14kg", "14kg /桶"),
        ("20KG", "20KG /桶"),
        ("无数字规格", "无数字规格"),
        ("20kg /桶", "20kg /桶"),
    ],
)
def test_format_specification_packaging_rules(spec, expected):
    assert format_specification_for_price_export(spec) == expected


@pytest.fixture
def client(monkeypatch, tmp_path):
    monkeypatch.delenv("FHD_API_KEYS", raising=False)
    monkeypatch.delenv("AUDIT_LOG_PATH", raising=False)
    # 默认读锁已开启：集成测试不显式带头时关闭门禁
    monkeypatch.setenv("FHD_DISABLE_DB_READ_LOCK", "1")
    with TestClient(http_app.app) as c:
        yield c


def test_health_ok_and_request_id_header(client):
    r = client.get("/api/health")
    assert r.status_code == 200
    assert r.json() == {"status": "ok"}
    assert r.headers.get("x-request-id")


def test_system_openapi_public(client):
    """GET /api/system/openapi：在仅反代 /api 的部署下提供 OpenAPI JSON。"""
    r = client.get("/api/system/openapi")
    assert r.status_code == 200
    data = r.json()
    assert "openapi" in data
    assert "paths" in data
    assert isinstance(data["paths"], dict)


def test_code_editor_status_public(client):
    r = client.get("/api/code-editor/status")
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j.get("phase") == "placeholder"


def test_code_editor_analyze_placeholder(client):
    r = client.post("/api/code-editor/analyze", json={"message": "x"})
    assert r.status_code == 200
    assert r.json().get("placeholder") is True


def test_document_templates_list_public(client):
    """GET /api/document-templates 公开；有 PG 与表时返回 success + data 数组。"""
    r = client.get("/api/document-templates", params={"role": "price_list_docx"})
    assert r.status_code == 200
    body = r.json()
    assert body.get("success") is True
    assert isinstance(body.get("data"), list)


def test_customers_import_endpoint_ok(client, monkeypatch):
    from io import BytesIO

    monkeypatch.setattr(
        "backend.customers_excel_import.run_customers_excel_import_bytes",
        lambda b: {
            "success": True,
            "inserted": 1,
            "updated": 1,
            "skipped_blank": 0,
            "synced_purchase_units": 2,
            "errors": [],
            "total_distinct_names": 2,
        },
    )
    monkeypatch.setattr("backend.routers.xcagi_compat._customers_write_raise", lambda: None)
    r = client.post(
        "/api/customers/import",
        files={
            "file": (
                "a.xlsx",
                BytesIO(b"x"),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j["data"]["inserted"] == 1
    assert j["data"]["updated"] == 1


def test_customers_import_endpoint_400_when_importer_fails(client, monkeypatch):
    from io import BytesIO

    monkeypatch.setattr(
        "backend.customers_excel_import.run_customers_excel_import_bytes",
        lambda b: {"success": False, "message": "工作表为空"},
    )
    monkeypatch.setattr("backend.routers.xcagi_compat._customers_write_raise", lambda: None)
    r = client.post(
        "/api/customers/import",
        files={"file": ("a.xlsx", BytesIO(b""), "application/octet-stream")},
    )
    assert r.status_code == 400
    assert "工作表为空" in (r.json().get("detail") or "")


def test_customers_match_empty_ok(client):
    r = client.get("/api/customers/match")
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j.get("matched") is None
    assert j.get("display") is None


def test_customers_match_uses_find_matching_customer(client, monkeypatch):
    monkeypatch.setattr("backend.shared_utils.extract_customer_name", lambda s: None)
    monkeypatch.setattr("backend.shared_utils.find_matching_customer", lambda s: "标准客户有限公司" if s else None)
    r = client.get("/api/customers/match", params={"customer_name": "某口语片段"})
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j.get("matched") == "标准客户有限公司"
    assert j.get("display") == "标准客户有限公司"


def test_fhd_identity_ok(client):
    r = client.get("/api/fhd/identity")
    assert r.status_code == 200
    data = r.json()
    assert data["backend"] == "fhd-http-app"
    assert data["xcagi_compat"] is True
    assert "routers" in data
    assert "backend.routers.xcagi_mods" in data["routers"]
    assert "system_industries" in data["endpoints"]
    assert "mods_list" in data["endpoints"]
    assert "startup_status" in data["endpoints"]
    assert "ai_unified_chat_stream" in data["endpoints"]


def test_unified_chat_stream_sse_stub(client, monkeypatch):
    """Planner 流式路由须存在；无真实 LLM 时用 mock 产出 SSE。"""
    monkeypatch.setattr(
        "backend.routers.xcagi_compat._ensure_vector_index_if_needed",
        lambda message, runtime_context: None,
    )

    def fake_sse(*args, **kwargs):
        yield {"type": "token", "text": "stub-stream"}

    monkeypatch.setattr("backend.routers.xcagi_compat.chat_stream_sse_events", fake_sse)
    r = client.post("/api/ai/unified_chat/stream", json={"message": "hello", "context": {}})
    assert r.status_code == 200
    assert "event-stream" in (r.headers.get("content-type") or "").lower()
    text = r.text
    assert "data:" in text
    assert "stub-stream" in text


def test_templates_api_pg_list_shape(client):
    r = client.get("/api/templates")
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert isinstance(j.get("templates"), list)
    assert isinstance(j.get("total"), int)
    r2 = client.get("/api/templates/list")
    assert r2.status_code == 200
    j2 = r2.json()
    assert j2.get("success") is True
    assert isinstance(j2.get("templates"), list)


def test_xcagi_stub_routes_no_404(client, monkeypatch):
    monkeypatch.setattr("backend.routers.xcagi_compat.run_agent_chat", lambda *a, **k: "stub-reply")
    assert client.get("/api/system/industries").status_code == 200
    assert client.get("/api/system/industry").status_code == 200
    assert client.post("/api/system/industry", json={}).status_code == 200
    ur = client.get("/api/products/units")
    assert ur.status_code == 200
    units_j = ur.json()
    assert units_j.get("success") is True
    assert units_j["data"][0]["name"] == "件"
    assert client.get("/api/products/list?page=1&per_page=100").status_code == 200
    assert client.get("/api/templates/list").status_code == 200
    assert client.get("/api/wechat_contacts/work_mode_feed?per_contact=1").status_code == 200
    client.post("/api/conversations/sessions/clear", json={"user_id": "default"})
    sm = client.post(
        "/api/conversations/message",
        json={"session_id": "s1", "user_id": "default", "role": "user", "content": "hi"},
    )
    assert sm.status_code == 200
    assert sm.json().get("success") is True
    ss = client.get("/api/conversations/sessions?limit=10&user_id=default")
    assert ss.status_code == 200
    assert ss.json().get("success") is True
    sess = ss.json().get("sessions") or []
    assert len(sess) == 1
    assert sess[0].get("session_id") == "s1"
    assert int(sess[0].get("message_count") or 0) == 1
    conv = client.get("/api/conversations/s1?user_id=default")
    assert conv.status_code == 200
    cm = conv.json().get("messages") or []
    assert len(cm) == 1 and cm[0].get("content") == "hi"
    assert client.post("/api/conversations/sessions/clear", json={"user_id": "default"}).status_code == 200
    assert (client.get("/api/conversations/sessions?user_id=default").json().get("sessions") or []) == []
    v2 = client.post("/api/ai/chat/v2", json={"message": "hello", "context": {}})
    assert v2.status_code == 200
    v2j = v2.json()
    assert v2j.get("success") is True
    assert v2j.get("response") == "stub-reply"
    assert client.post("/api/ai/conversation/new", json={}).status_code == 200
    assert client.get("/api/ai/context").status_code == 200
    assert client.post("/api/ai/context/clear", json={}).status_code == 200
    assert client.get("/api/ai/config").status_code == 200
    tts = client.post("/api/tts/synthesize", json={"text": "你好", "lang": "zh"})
    assert tts.status_code == 200
    assert tts.json().get("success") is False
    r = client.get("/api/conversations/mnobzv9gcpp65pidzxb")
    assert r.status_code == 200
    assert r.json()["id"] == "mnobzv9gcpp65pidzxb"
    assert r.json()["messages"] == []
    pu = client.get("/api/purchase_units")
    assert pu.status_code == 200
    assert pu.json().get("success") is True
    assert "data" in pu.json()
    assert isinstance(pu.json()["data"], list)
    cr = client.get("/api/customers")
    assert cr.status_code == 200
    assert isinstance(cr.json().get("data"), list)
    clr = client.get("/api/customers/list?page=1&per_page=20")
    assert clr.status_code == 200
    assert isinstance(clr.json().get("data"), list)
    assert isinstance(clr.json().get("total"), int)
    assert client.get("/api/system/test-db/status").status_code == 200
    assert client.get("/api/preferences?user_id=default").status_code == 200
    assert client.get("/api/distillation/versions").status_code == 200
    assert client.get("/api/intent-packages").status_code == 200
    tr = client.get("/api/tools")
    assert tr.status_code == 200
    tj = tr.json()
    assert tj.get("success") is True
    assert isinstance(tj.get("tools"), list) and len(tj["tools"]) > 0
    assert client.get("/api/db-tools").status_code == 200
    assert client.get("/api/db-tools").json().get("tools") == tj["tools"]
    tc = client.get("/api/tool-categories")
    assert tc.status_code == 200
    assert tc.json().get("success") is True
    assert isinstance(tc.json().get("categories"), list)
    assert client.post("/api/preferences", json={"theme": "dark"}).status_code == 200


@_skip_legacy_sqlite_file_tests
def test_purchase_units_reads_sqlite(client, monkeypatch, tmp_path):
    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE purchase_units (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            unit_name TEXT NOT NULL,
            contact_person TEXT,
            contact_phone TEXT,
            address TEXT,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO purchase_units (unit_name, is_active) VALUES (?, 1)",
        ("测试购买单位",),
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    r = client.get("/api/purchase_units")
    assert r.status_code == 200
    data = r.json()["data"]
    assert len(data) == 1
    assert data[0]["unit_name"] == "测试购买单位"
    assert data[0]["is_active"] == 1
    r_units = client.get("/api/products/units")
    assert r_units.status_code == 200
    uj = r_units.json()
    assert uj.get("success") is True
    assert len(uj["data"]) == 1
    assert uj["data"][0]["name"] == "测试购买单位"

    cr = client.get("/api/customers")
    assert cr.status_code == 200
    cdata = cr.json()["data"]
    assert len(cdata) == 1
    assert cdata[0]["customer_name"] == "测试购买单位"
    assert cdata[0]["id"] == 1
    lr = client.get("/api/customers/list?page=1&per_page=10")
    assert lr.status_code == 200
    assert lr.json()["total"] == 1
    assert len(lr.json()["data"]) == 1


@_skip_legacy_sqlite_file_tests
def test_customers_falls_back_to_customers_db_when_products_empty(client, monkeypatch, tmp_path):
    """主库 customers 表无数据时，读独立 customers.db 的 purchase_units。"""
    main = tmp_path / "products.db"
    conn = sqlite3.connect(str(main))
    conn.execute(
        """
        CREATE TABLE customers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            customer_name TEXT NOT NULL,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn.commit()
    conn.close()

    side = tmp_path / "customers.db"
    conn2 = sqlite3.connect(str(side))
    conn2.execute(
        """
        CREATE TABLE purchase_units (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            unit_name TEXT NOT NULL,
            contact_person TEXT,
            contact_phone TEXT,
            address TEXT,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn2.execute(
        "INSERT INTO purchase_units (unit_name, is_active) VALUES (?, 1), (?, 1)",
        ("客户甲", "客户乙"),
    )
    conn2.commit()
    conn2.close()

    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: main)
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_customers_db_path", lambda: side)

    r = client.get("/api/customers")
    assert r.status_code == 200
    names = sorted(c["customer_name"] for c in r.json()["data"])
    assert names == ["客户乙", "客户甲"]
    lr = client.get("/api/customers/list?page=1&per_page=10")
    assert lr.status_code == 200
    assert lr.json()["total"] == 2


@_skip_legacy_sqlite_file_tests
def test_customers_aligns_with_products_units_when_only_products_unit_column(
    client, monkeypatch, tmp_path,
):
    """客户名仅出现在 products.unit、无 purchase_units/customers 表时，客户总数与单位下拉一致。"""
    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            model_number TEXT NOT NULL,
            name TEXT NOT NULL,
            unit TEXT,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO products (model_number, name, unit) VALUES (?, ?, ?), (?, ?, ?)",
        ("m1", "n1", "半岛风情", "m2", "n2", "奔奔熊鞋柜"),
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_customers_db_path", lambda: None)

    units = client.get("/api/products/units").json()["data"]
    cust = client.get("/api/customers").json()["data"]
    assert len(cust) == len(units) == 2
    assert {c["customer_name"] for c in cust} == {u["name"] for u in units}
    lr = client.get("/api/customers/list?page=1&per_page=20")
    assert lr.json()["total"] == 2


@_skip_legacy_sqlite_file_tests
def test_purchase_units_merge_includes_products_unit_column(client, monkeypatch, tmp_path):
    """主表已有购买单位时，仍应出现仅写在 products.unit 上的新单位（与 XCAGI 下拉逻辑一致）。"""
    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE purchase_units (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            unit_name TEXT NOT NULL,
            contact_person TEXT,
            contact_phone TEXT,
            address TEXT,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO purchase_units (unit_name, is_active) VALUES (?, 1)",
        ("已有客户A",),
    )
    conn.execute(
        """
        CREATE TABLE products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            model_number TEXT,
            name TEXT,
            specification TEXT,
            price REAL,
            quantity INTEGER,
            unit TEXT,
            is_active INTEGER DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
        "VALUES ('M1','一号','',1,1,'新单位仅在产品表',1)"
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    pu = client.get("/api/purchase_units")
    assert pu.status_code == 200
    body = pu.json()
    assert body.get("success") is True
    names = {row["unit_name"] for row in body["data"]}
    assert names == {"已有客户A", "新单位仅在产品表"}
    ur = client.get("/api/products/units")
    assert ur.status_code == 200
    unames = {x["name"] for x in ur.json()["data"]}
    assert unames == {"已有客户A", "新单位仅在产品表"}


@_skip_legacy_sqlite_file_tests
def test_products_list_reads_sqlite_and_matches_frontend_shape(client, monkeypatch, tmp_path):
    """GET /api/products/list 返回 { success, data: 行数组, total }，供 Pinia 直接使用。"""
    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            model_number TEXT,
            name TEXT,
            specification TEXT,
            price REAL,
            quantity INTEGER,
            unit TEXT,
            is_active INTEGER DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
        "VALUES ('M1','一号','规格A',10.5,1,'件',1)"
    )
    conn.execute(
        "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
        "VALUES ('M2','二号','',2,0,'千克',1)"
    )
    conn.execute(
        "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
        "VALUES ('X','隐藏','',1,0,'件',0)"
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)

    r = client.get("/api/products/list?page=1&per_page=10")
    assert r.status_code == 200
    body = r.json()
    assert body["success"] is True
    assert isinstance(body["data"], list)
    assert body["total"] == 2
    assert len(body["data"]) == 2

    r_kw = client.get("/api/products/list?page=1&per_page=10&keyword=二")
    assert r_kw.json()["total"] == 1
    assert r_kw.json()["data"][0]["model_number"] == "M2"

    r_u = client.get("/api/products/list", params={"page": 1, "per_page": 10, "unit": "件"})
    assert r_u.json()["total"] == 1
    assert r_u.json()["data"][0]["unit"] == "件"


@_skip_legacy_sqlite_file_tests
def test_products_export_docx_fills_unit_date_and_table(client, monkeypatch, tmp_path):
    tpl = tmp_path / "mini_price_tpl.docx"
    d = Document()
    d.add_paragraph("客户名称：\t报价日期：")
    t = d.add_table(rows=3, cols=6)
    t.rows[0].cells[5].text = "备注"
    d.save(str(tpl))

    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            model_number TEXT,
            name TEXT,
            specification TEXT,
            price REAL,
            quantity INTEGER,
            unit TEXT,
            is_active INTEGER DEFAULT 1
        )
        """
    )
    conn.execute(
        "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
        "VALUES ('A1','产品甲','5kg',12.3,1,'客户甲',1)"
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    monkeypatch.setattr(
        "backend.routers.xcagi_compat.resolve_price_list_docx_template",
        lambda *_a, **_k: tpl,
    )

    r = client.get(
        "/api/products/price-list-export",
        params={"unit": "客户甲", "date": "2026-04-11"},
    )
    assert r.status_code == 200
    assert "wordprocessingml.document" in (r.headers.get("content-type") or "")

    out = Document(io.BytesIO(r.content))
    para_hit = False
    for p in out.paragraphs:
        if "客户甲" in p.text and "2026-04-11" in p.text:
            para_hit = True
            break
    assert para_hit

    tbl = out.tables[0]
    assert tbl.rows[1].cells[0].text.strip() == "A1"
    assert tbl.rows[1].cells[1].text.strip() == "产品甲"
    assert tbl.rows[1].cells[2].text.strip() == "5kg /缶"
    assert tbl.rows[1].cells[3].text.strip() == "12.30 元/kg"

    r_legacy = client.get(
        "/api/products/export.docx",
        params={"unit": "客户甲", "date": "2026-04-11"},
    )
    assert r_legacy.status_code == 200


def test_products_price_list_template_preview(client, monkeypatch, tmp_path):
    tpl = tmp_path / "preview_price_tpl.docx"
    d = Document()
    t = d.add_table(rows=2, cols=4)
    t.rows[0].cells[0].text = "编号"
    t.rows[0].cells[1].text = "产品名称"
    t.rows[0].cells[2].text = "规格"
    t.rows[0].cells[3].text = "单价"
    t.rows[1].cells[0].text = "X1"
    t.rows[1].cells[1].text = "样例"
    t.rows[1].cells[2].text = "1kg"
    t.rows[1].cells[3].text = "9.9"
    d.save(str(tpl))

    monkeypatch.setattr(
        "backend.price_list_docx_export.resolve_price_list_docx_template",
        lambda *_a, **_k: tpl,
    )
    r = client.get("/api/products/price-list-template-preview")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    assert data["headers"] == ["编号", "产品名称", "规格", "单价"]
    assert data["sample_rows"][0]["产品名称"] == "样例"
    assert "template_hint" in data


@_skip_legacy_sqlite_file_tests
def test_products_export_docx_adds_rows_when_table_too_short(client, monkeypatch, tmp_path):
    """模板表格行数少于产品条数时，应自动 add_row 再写入。"""
    tpl = tmp_path / "mini_price_short_table.docx"
    d = Document()
    d.add_paragraph("客户名称：\t报价日期：")
    t = d.add_table(rows=2, cols=6)
    t.rows[0].cells[5].text = "备注"
    d.save(str(tpl))

    db = tmp_path / "products_many.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE products (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            model_number TEXT,
            name TEXT,
            specification TEXT,
            price REAL,
            quantity INTEGER,
            unit TEXT,
            is_active INTEGER DEFAULT 1
        )
        """
    )
    for idx in range(4):
        conn.execute(
            "INSERT INTO products (model_number,name,specification,price,quantity,unit,is_active) "
            "VALUES (?,?,?,?,1,'客户甲',1)",
            (f"M{idx}", f"产品{idx}", "1kg", float(idx + 1)),
        )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    monkeypatch.setattr(
        "backend.routers.xcagi_compat.resolve_price_list_docx_template",
        lambda *_a, **_k: tpl,
    )

    r = client.get("/api/products/price-list-export", params={"unit": "客户甲"})
    assert r.status_code == 200
    out = Document(io.BytesIO(r.content))
    tbl = out.tables[0]
    assert len(tbl.rows) >= 1 + 4
    model_cols = [tbl.rows[r].cells[0].text.strip() for r in range(1, 5)]
    assert set(model_cols) == {"M0", "M1", "M2", "M3"}
    # 价格 4.00 元/kg 应对应 id 最大的一条（M3）
    assert any("4.00 元/kg" in tbl.rows[r].cells[3].text for r in range(1, len(tbl.rows)))


def test_xcagi_stubs_exempt_from_api_key(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r = client.get("/api/system/industries")
    assert r.status_code == 200
    r_dt = client.get("/api/document-templates")
    assert r_dt.status_code == 200
    r2 = client.get("/api/products/list")
    assert r2.status_code == 200
    r3 = client.get("/api/conversations/any-id")
    assert r3.status_code == 200
    r4 = client.get("/api/purchase_units")
    assert r4.status_code == 200
    r4b = client.get("/api/customers/list?page=1&per_page=5")
    assert r4b.status_code == 200
    r5 = client.get("/api/system/test-db/status")
    assert r5.status_code == 200
    r6 = client.get("/api/preferences")
    assert r6.status_code == 200


def test_client_mods_off_trailing_slash_ok(client):
    r = client.post("/api/state/client-mods-off/", json={"client_mods_off": True})
    assert r.status_code == 200
    assert r.json()["client_mods_off"] is True


def test_health_preserves_incoming_request_id(client):
    rid = "test-req-id-001"
    r = client.get("/api/health", headers={"X-Request-ID": rid})
    assert r.status_code == 200
    assert r.headers.get("x-request-id") == rid


def test_audit_log_jsonlines(client, monkeypatch, tmp_path):
    log_path = tmp_path / "audit.jsonl"
    monkeypatch.setenv("AUDIT_LOG_PATH", str(log_path))
    r = client.get("/api/health")
    assert r.status_code == 200
    text = log_path.read_text(encoding="utf-8").strip()
    assert text
    row = json.loads(text.splitlines()[-1])
    assert row["path"] == "/api/health"
    assert row["method"] == "GET"
    assert row["status_code"] == 200
    assert row["request_id"] == r.headers.get("x-request-id")


def test_api_key_blocks_api_without_key(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1,k2")
    r = client.post("/api/chat", json={"message": "hello"})
    assert r.status_code == 401
    assert r.json().get("detail")


def test_api_key_allows_with_header(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    monkeypatch.setattr(http_app, "chat", lambda *a, **k: "stub")
    r = client.post(
        "/api/chat",
        json={"message": "hello"},
        headers={"X-API-Key": "k1"},
    )
    assert r.status_code == 200
    assert r.json() == {"reply": "stub"}


def test_health_exempt_from_api_key(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r = client.get("/api/health")
    assert r.status_code == 200


def test_system_test_db_disable_route_exists(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r = client.post("/api/system/test-db/disable", json={})
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j.get("data", {}).get("test_db_enabled") is False


def test_mods_list_ok_and_exempt_from_api_key(client, monkeypatch):
    r = client.get("/api/mods/")
    assert r.status_code == 200
    payload = r.json()
    assert payload.get("success") is True
    data = payload["data"]
    assert isinstance(data, list) and len(data) >= 1
    ids = [x.get("id") for x in data]
    assert "all" in ids  # 内置目录项仍合并；真实扩展在 XCAGI/mods 时排在前面
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r2 = client.get("/api/mods")
    assert r2.status_code == 200
    p2 = r2.json()
    assert len(p2["data"]) == len(data)


def test_mods_loading_status_and_routes_ok(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    ls = client.get("/api/mods/loading-status")
    assert ls.status_code == 200
    assert ls.json().get("success") is True
    assert ls.json().get("loaded") is True
    rt = client.get("/api/mods/routes")
    assert rt.status_code == 200
    body = rt.json()
    assert body.get("success") is True
    assert isinstance(body.get("data"), list)


def test_startup_status_ok_and_exempt_from_api_key(client, monkeypatch):
    r = client.get("/api/startup/status")
    assert r.status_code == 200
    body = r.json()
    assert body.get("ready") is True
    assert body.get("progress_percent") == 100
    comps = body.get("components") or []
    assert any(c.get("name") == "mods" and c.get("status") == "ready" for c in comps)
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r2 = client.get("/api/startup/status/")
    assert r2.status_code == 200
    assert r2.json().get("ready") is True


def test_shipment_records_units_matches_products_units(client, monkeypatch):
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    pu = client.get("/api/products/units")
    su = client.get("/api/shipment/shipment-records/units")
    assert su.status_code == 200
    assert su.json() == pu.json()


def test_client_mods_off_ok(client):
    r = client.post("/api/state/client-mods-off", json={"client_mods_off": True})
    assert r.status_code == 200
    assert r.json() == {"success": True, "client_mods_off": True}


def test_products_list_empty_when_x_client_mods_off_with_mod_gate(client, monkeypatch):
    """原版模式请求头：在扩展 Mod 门禁开启时也应隐藏业务读（与 POST client-mods-off 是否持久化无关）。"""
    monkeypatch.setenv("FHD_BUSINESS_DATA_REQUIRES_EXTENSION_MOD", "1")
    monkeypatch.setattr(
        "backend.shell.mod_business_scope.extension_mod_manifest_rows",
        lambda: [{"id": "ext1", "type": "mod"}],
    )
    r0 = client.get("/api/products/list?page=1&per_page=5")
    assert r0.status_code == 200
    assert r0.json().get("success") is True
    r1 = client.get(
        "/api/products/list?page=1&per_page=5",
        headers={"X-Client-Mods-Off": "1"},
    )
    assert r1.status_code == 200
    j1 = r1.json()
    assert j1.get("success") is True
    assert j1.get("total") == 0
    assert j1.get("data") == []
    hint = str(j1.get("schema_hint") or "")
    assert "原版模式" in hint or "X-Client-Mods-Off" in hint


def test_products_list_empty_when_x_client_mods_off_even_if_mod_gate_env_off(client, monkeypatch):
    """未启用 FHD_BUSINESS_DATA_REQUIRES_EXTENSION_MOD 时，仅凭原版模式头也应隐藏产品列表。"""
    monkeypatch.delenv("FHD_BUSINESS_DATA_REQUIRES_EXTENSION_MOD", raising=False)
    r1 = client.get(
        "/api/products/list?page=1&per_page=5",
        headers={"X-Client-Mods-Off": "1"},
    )
    assert r1.status_code == 200
    j1 = r1.json()
    assert j1.get("success") is True
    assert j1.get("total") == 0
    assert j1.get("data") == []
    assert "原版模式" in str(j1.get("schema_hint") or "")


def test_traditional_mode_list_root(client, monkeypatch, tmp_path):
    monkeypatch.setenv("WORKSPACE_ROOT", str(tmp_path))
    r = client.get("/api/traditional-mode/list")
    assert r.status_code == 200
    data = r.json()
    assert data["success"] is True
    assert data["path"] == ""
    assert data["items"] == []


def test_traditional_mode_watch_route_registered():
    """SSE 长连接在 TestClient 下整块读取易阻塞，只校验路由存在。"""
    paths = []
    for route in http_app.app.routes:
        p = getattr(route, "path", None)
        if p:
            paths.append(p)
    assert "/api/traditional-mode/watch" in paths


def test_state_and_traditional_exempt_from_api_key(client, monkeypatch, tmp_path):
    monkeypatch.setenv("WORKSPACE_ROOT", str(tmp_path))
    monkeypatch.setenv("FHD_API_KEYS", "k1")
    r = client.post("/api/state/client-mods-off", json={"client_mods_off": False})
    assert r.status_code == 200
    r2 = client.get("/api/traditional-mode/list")
    assert r2.status_code == 200
    rid = client.get("/api/fhd/identity")
    assert rid.status_code == 200


def test_chat_accepts_legacy_body_aliases(client, monkeypatch):
    captured: dict = {}

    def fake_chat(message, runtime_context=None, system_prompt=None):
        captured["message"] = message
        captured["runtime_context"] = runtime_context
        captured["system_prompt"] = system_prompt
        return "ok"

    monkeypatch.setattr(http_app, "chat", fake_chat)
    r = client.post(
        "/api/chat",
        json={
            "content": "hello",
            "neuro_context": {"excel_file_path": "uploads/x.xlsx"},
            "instructions": "sys",
            "llm_mode": "offline",
        },
    )
    assert r.status_code == 200
    assert r.json() == {"reply": "ok"}
    assert captured["message"] == "hello"
    assert captured["runtime_context"] == {"excel_file_path": "uploads/x.xlsx"}
    assert captured["system_prompt"] == "sys"


def test_xcagi_chat_extracts_excel_path_into_runtime_context(client, monkeypatch):
    captured: dict = {}

    def fake_run_agent_chat(message, runtime_context=None, system_prompt=None):
        captured["message"] = message
        captured["runtime_context"] = runtime_context
        captured["system_prompt"] = system_prompt
        return "ok"

    monkeypatch.setattr("backend.routers.xcagi_compat.run_agent_chat", fake_run_agent_chat)
    r = client.post(
        "/api/ai/chat",
        json={"message": "请读取 @424/26年出货单打印/鸿瑞达报价26年.xlsx 并分析"},
    )
    assert r.status_code == 200
    payload = r.json()
    assert payload.get("success") is True
    assert captured["runtime_context"]["excel_file_path"] == "424/26年出货单打印/鸿瑞达报价26年.xlsx"
    assert "424/26年出货单打印/鸿瑞达报价26年.xlsx" in captured["runtime_context"]["excel_file_paths"]


def test_xcagi_chat_prefers_context_file_path_when_message_has_only_basename(client, monkeypatch):
    captured: dict = {}

    def fake_run_agent_chat(message, runtime_context=None, system_prompt=None):
        captured["message"] = message
        captured["runtime_context"] = runtime_context
        captured["system_prompt"] = system_prompt
        return "ok"

    monkeypatch.setattr("backend.routers.xcagi_compat.run_agent_chat", fake_run_agent_chat)
    r = client.post(
        "/api/ai/chat",
        json={
            "message": "@鸿瑞达报价26年.xlsx 添加新的购买单位和对应产品",
            "context": {
                "excel_analysis": {
                    "preview_data": {
                        "file_path": "uploads/abc123_鸿瑞达报价26年.xlsx"
                    }
                }
            },
        },
    )
    assert r.status_code == 200
    payload = r.json()
    assert payload.get("success") is True
    assert captured["runtime_context"]["excel_file_path"] == "uploads/abc123_鸿瑞达报价26年.xlsx"
    assert captured["runtime_context"]["excel_file_paths"][0] == "uploads/abc123_鸿瑞达报价26年.xlsx"


def test_xcagi_chat_interrupt_returns_cleared_runtime_context(client, monkeypatch):
    calls: list[int] = []

    def fail_if_called(*_a, **_k):
        calls.append(1)
        raise AssertionError("run_agent_chat should not run on reflex interrupt")

    monkeypatch.setattr("backend.routers.xcagi_compat.run_agent_chat", fail_if_called)
    r = client.post(
        "/api/ai/chat",
        json={
            "message": "取消",
            "context": {
                "excel_file_path": "uploads/x.xlsx",
                "recent_messages": [{"role": "user", "content": "old"}],
                "keep": True,
            },
        },
    )
    assert r.status_code == 200
    assert not calls
    payload = r.json()
    assert payload.get("success") is True
    data = payload.get("data") or {}
    assert "好的" in str(data.get("text") or payload.get("response") or "")
    assert data.get("runtime_context") == {"keep": True}


def test_xcagi_chat_vector_request_returns_actionable_error(client, monkeypatch):
    called = {"chat": False}

    def fake_run_agent_chat(message, runtime_context=None, system_prompt=None):
        _ = (message, runtime_context, system_prompt)
        called["chat"] = True
        return "should_not_happen"

    monkeypatch.setattr("backend.routers.xcagi_compat.run_agent_chat", fake_run_agent_chat)
    monkeypatch.setattr(
        "backend.routers.xcagi_compat.execute_workflow_tool",
        lambda *args, **kwargs: '{"error":"file_not_found","message":"missing"}',
    )
    r = client.post(
        "/api/ai/chat",
        json={"message": "请给 @uploads/missing.xlsx 建立向量索引"},
    )
    assert r.status_code == 200
    payload = r.json()
    assert payload.get("success") is True
    assert "建立向量索引失败" in payload.get("response", "")
    assert called["chat"] is False


def test_chat_503_when_llm_not_configured(client, monkeypatch):
    def fail_key():
        raise RuntimeError("未配置大模型 API Key：请设置 DP_API_KEY")

    # planner 在 import 时绑定了 require_api_key，需 patch 其模块内引用
    monkeypatch.setattr("backend.planner.require_api_key", fail_key)
    r = client.post("/api/chat", json={"message": "hi"})
    assert r.status_code == 503
    body = r.json()
    assert "detail" in body
    assert "未配置" in body["detail"]


def test_upload_excel_returns_legacy_path_and_context(client, monkeypatch, tmp_path):
    monkeypatch.setenv("WORKSPACE_ROOT", str(tmp_path))
    from io import BytesIO

    xlsx = tmp_path / "dummy.xlsx"
    xlsx.write_bytes(b"dummy")
    r = client.post(
        "/api/upload/excel",
        files={"file": ("dummy.xlsx", BytesIO(xlsx.read_bytes()), "application/octet-stream")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["file_path"] == data["path"]
    assert data["runtime_context"] == data["context"]
    assert "excel_file_path" in data["runtime_context"]


@_skip_legacy_sqlite_file_tests
def test_customers_crud_writes_purchase_units(client, monkeypatch, tmp_path):
    db = tmp_path / "products.db"
    conn = sqlite3.connect(str(db))
    conn.execute(
        """
        CREATE TABLE purchase_units (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            unit_name TEXT NOT NULL,
            contact_person TEXT,
            contact_phone TEXT,
            address TEXT,
            is_active INTEGER NOT NULL DEFAULT 1
        )
        """
    )
    conn.execute(
        "CREATE TABLE products (id INTEGER PRIMARY KEY, model_number TEXT, name TEXT, unit TEXT)"
    )
    conn.execute(
        "INSERT INTO products (model_number, name, unit) VALUES (?, ?, ?)",
        ("m1", "产品1", "公司乙"),
    )
    conn.commit()
    conn.close()
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_products_db_path", lambda: db)
    monkeypatch.setattr("backend.routers.xcagi_compat.resolve_customers_db_path", lambda: None)

    r = client.post(
        "/api/customers",
        json={"customer_name": "公司甲", "contact_person": "张三"},
    )
    assert r.status_code == 200
    j = r.json()
    assert j.get("success") is True
    assert j["data"]["customer_name"] == "公司甲"
    cid = j["data"]["id"]

    assert client.get(f"/api/customers/{cid}").status_code == 200
    r2 = client.put(
        f"/api/customers/{cid}",
        json={"customer_name": "公司乙", "contact_person": "李四"},
    )
    assert r2.status_code == 200
    assert r2.json()["data"]["customer_name"] == "公司乙"

    assert client.delete(f"/api/customers/{cid}").status_code == 200
    conn2 = sqlite3.connect(str(db))
    cur2 = conn2.cursor()
    assert cur2.execute("SELECT COUNT(*) FROM purchase_units").fetchone()[0] == 0
    assert cur2.execute("SELECT COUNT(*) FROM products").fetchone()[0] == 0
    conn2.close()


def test_products_write_api_add_update_delete_roundtrip(client, monkeypatch):
    """POST /api/products/add、/update、/delete 与 XCAGI 产品页「保存」一致（须二级写令牌）。"""
    import uuid

    from sqlalchemy import inspect

    from backend.database import get_sync_engine

    try:
        eng = get_sync_engine()
        insp = inspect(eng)
    except Exception as e:
        pytest.skip(f"数据库不可用: {e}")
    if "products" not in insp.get_table_names():
        pytest.skip("public.products 表不存在")

    monkeypatch.setenv("FHD_DB_WRITE_TOKEN", "pytest_products_write")
    w_headers = {"X-FHD-Db-Write-Token": "pytest_products_write"}

    tag = f"zzapi_{uuid.uuid4().hex[:12]}"
    slug_unit = f"单元测试_{tag[:8]}"

    r = client.post(
        "/api/products/add",
        headers=w_headers,
        json={
            "model_number": tag,
            "name": "接口测产品",
            "unit": slug_unit,
            "specification": "4KG/卡",
            "price": 10.0,
        },
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body.get("success") is True
    pid = body["data"]["id"]
    assert isinstance(pid, int)

    r_up = client.post(
        "/api/products/update",
        headers=w_headers,
        json={
            "id": pid,
            "model_number": tag,
            "name": "接口测产品已改",
            "unit": slug_unit,
            "specification": "4KG/卡",
            "price": 77.77,
        },
    )
    assert r_up.status_code == 200, r_up.text
    assert r_up.json().get("success") is True

    lst = client.get("/api/products/list", params={"keyword": tag, "per_page": 50})
    assert lst.status_code == 200
    rows = lst.json().get("data") or []
    hit = [x for x in rows if str(x.get("model_number") or "") == tag]
    assert len(hit) >= 1
    pr = float(hit[0].get("price") or 0)
    assert abs(pr - 77.77) < 0.02

    r_del = client.post("/api/products/delete", headers=w_headers, json={"id": pid})
    assert r_del.status_code == 200, r_del.text
    assert r_del.json().get("success") is True

    lst2 = client.get("/api/products/list", params={"keyword": tag, "per_page": 50})
    hit2 = [x for x in (lst2.json().get("data") or []) if str(x.get("model_number") or "") == tag]
    assert not hit2
