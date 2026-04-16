"""
使用仓库内 `424/document_templates/price_list_default.docx`（或环境变量指定路径）生成报价 Word：
填写「客户名称 / 报价日期」行，并在首个表格中写入产品行（编号、名称、规格、单价，单价列带单位「元/kg」）。
规格列按千克数追加「/卡、/缶、/桶」：≤4kg 为「/卡」；4kg 以上且不足 14kg 为「/缶」（含 10kg 及以下与 10～14kg 段）；≥14kg 为「/桶」。
"""

from __future__ import annotations

import io
import logging
import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger(__name__)


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[1]


def resolve_price_list_docx_template(template_slug: str | None = None) -> Path | None:
    """
    优先 PostgreSQL ``document_templates``（``role=price_list_docx``），
    失败时见 ``backend.document_template_service`` 内置降级（含 ``FHD_PRICE_LIST_DOCX_TEMPLATE``）。
    """
    from backend.document_template_service import ROLE_PRICE_LIST, resolve_template_path

    return resolve_template_path(role=ROLE_PRICE_LIST, slug=(template_slug or "").strip() or None)


def _format_price(value: Any) -> str:
    try:
        if value is None:
            return ""
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return str(value or "")


_PRICE_UNIT_SUFFIX = "元/kg"


def _parse_kg_from_specification(text: str) -> float | None:
    """
    从规格文案中解析千克数（如 ``5kg``、``10 公斤``、``20KG``）。
    无法解析时返回 None（保持原文不写规则单位）。
    """
    s = (text or "").strip()
    if not s:
        return None
    u = (
        s.upper()
        .replace("ＫＧ", "KG")
        .replace("千克", "KG")
        .replace("公斤", "KG")
        .replace("kｇ", "KG")
    )
    m = re.search(r"(\d+(?:\.\d+)?)\s*(?:KG|K)", u, re.I)
    if m:
        return float(m.group(1))
    m2 = re.search(r"(\d+(?:\.\d+)?)\s*(?:千克|公斤)", s)
    if m2:
        return float(m2.group(1))
    m3 = re.search(r"(\d+(?:\.\d+)?)", s)
    if m3:
        return float(m3.group(1))
    return None


def _spec_already_has_packaging_suffix(base: str) -> bool:
    b = (base or "").rstrip()
    return bool(re.search(r"(/|／)(卡|缶|桶)\s*$", b))


def format_specification_for_price_export(raw: Any) -> str:
    """
    导出用规格展示：在原文后追加包装档位（与数据库字段无关，仅影响 Word）。

    规则（以解析到的千克数为准，与业务约定一致）：
    - **4kg 及其以下** → 追加 ``/卡``
    - **大于 4kg 且小于 14kg**（含「10 公斤及其下」与 10～14kg 之间）→ 追加 ``/缶``
    - **14kg 及以上** → 追加 ``/桶``

    解析不到千克数时保持原文（不追加）。
    """
    base = str(raw or "").strip()
    if not base:
        return ""
    if _spec_already_has_packaging_suffix(base):
        return base
    kg = _parse_kg_from_specification(base)
    if kg is None:
        return base
    if kg <= 4:
        suffix = "/卡"
    elif kg < 14:
        suffix = "/缶"
    else:
        suffix = "/桶"
    return f"{base} {suffix}"


def _format_price_with_unit(value: Any) -> str:
    """单价格展示为「数值 + 元/kg」；空值不写单位；已含单位则不重复。"""
    raw = _format_price(value)
    t = (raw or "").strip()
    if not t:
        return ""
    if "元/kg" in t or "元／kg" in t:
        return t
    return f"{t} {_PRICE_UNIT_SUFFIX}"


def _strip_vertical_merge_from_row(tr: Any) -> None:
    """
    去掉行内各单元格的纵向合并标记。
    含 ``w:vMerge`` 的「延续」行在 python-docx 里会指回上一行同一格，写入多条产品会叠在同格；
    克隆模板行后必须先清掉再追加。
    """
    for tc in tr.tc_lst:
        tc_pr = tc.tcPr
        if tc_pr is None:
            continue
        vm = tc_pr.find(qn("w:vMerge"))
        if vm is not None:
            tc_pr.remove(vm)


def _physical_row_count(table: Table) -> int:
    return len(table._tbl.tr_lst)


def _append_table_row_clone(table: Table, template_row_index: int) -> None:
    """
    在表末追加一行：深拷贝 ``template_row_index`` 对应 ``w:tr``（保留边框/格宽），
    并清除纵向合并；失败时退回 ``Table.add_row()``。
    """
    n = _physical_row_count(table)
    if template_row_index < 0 or template_row_index >= n:
        table.add_row()
        return
    try:
        src_tr = table.rows[template_row_index]._tr
        new_tr = deepcopy(src_tr)
        _strip_vertical_merge_from_row(new_tr)
        table._tbl.append(new_tr)
    except Exception:
        logger.exception("克隆表格行失败，改用 add_row")
        table.add_row()


def _paragraph_clear_body(paragraph: Paragraph) -> None:
    """保留 ``w:pPr``，删除正文（run / 超链接等），便于重写带格式的 runs。"""
    p = paragraph._p
    for el in list(p):
        if el.tag not in (qn("w:pPr"),):
            p.remove(el)


def _fill_customer_and_quote_date(doc: Document, customer: str, quote_date: str) -> bool:
    """在包含「客户名称」「报价日期」的段落中写入内容；**客户名称与报价日期数值**加单下划线。返回是否已替换。"""
    customer = (customer or "").strip() or "（未指定购买单位）"
    quote_date = (quote_date or "").strip()
    for para in doc.paragraphs:
        text = para.text or ""
        if "客户名称" in text and "报价日期" in text:
            _paragraph_clear_body(para)
            para.add_run("客户名称：")
            r_cust = para.add_run(customer)
            r_cust.font.underline = WD_UNDERLINE.SINGLE
            para.add_run("\t")
            para.add_run("报价日期：")
            r_date = para.add_run(quote_date)
            r_date.font.underline = WD_UNDERLINE.SINGLE
            return True
    return False


def _fill_first_table_products(doc: Document, products: list[dict[str, Any]], start_row: int = 1) -> int:
    """
    将产品写入文档中第一个表格，从 ``start_row`` 行开始（0 起始，默认第 2 行为首条数据）。
    列：0 型号/编号、1 名称、2 规格（带包装档位）、3 单价（带「元/kg」）；第 5 列保留为备注（不覆盖）。
    行数不足时在表末追加行：优先 **克隆首条数据行样式**（并去掉 ``w:vMerge``），
    含纵向合并的合同模板下 ``add_row()`` 往往无法形成独立格子；克隆失败再 ``add_row()``。
    返回写入行数。
    """
    if not doc.tables:
        return 0
    table = doc.tables[0]
    if start_row < 0:
        start_row = 0

    clone_src = start_row if start_row < _physical_row_count(table) else max(0, _physical_row_count(table) - 1)

    written = 0
    for i, row in enumerate(products):
        ri = start_row + i
        while _physical_row_count(table) <= ri:
            _append_table_row_clone(table, clone_src)
        cells = table.rows[ri].cells
        if len(cells) >= 1:
            cells[0].text = str(row.get("model_number") or "").strip()
        if len(cells) >= 2:
            cells[1].text = str(row.get("name") or "").strip()
        if len(cells) >= 3:
            cells[2].text = format_specification_for_price_export(row.get("specification"))
        if len(cells) >= 4:
            cells[3].text = _format_price_with_unit(row.get("price"))
        written += 1

    # 清空后续数据行前四列，避免旧数据残留
    for ri in range(start_row + len(products), len(table.rows)):
        cells = table.rows[ri].cells
        for j in range(min(4, len(cells))):
            cells[j].text = ""

    return written


def _center_table_cell_text(table: Table) -> None:
    """表格内文字水平、垂直居中（导出数据与模板原有单元格一并处理）。"""
    for row in table.rows:
        for cell in row.cells:
            try:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            except Exception:
                pass
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _table_row_texts(row: Any) -> list[str]:
    out: list[str] = []
    for cell in row.cells:
        out.append((cell.text or "").strip())
    return out


def read_docx_first_table_preview(
    template_path: Path, *, sheet_name: str = "Word（首表）"
) -> dict[str, Any]:
    """
    读取任意 ``.docx`` 中**首个表格**的第 1 行表头与第 2 行示例（若存在）。
    价表与销售合同预览共用。
    """
    try:
        doc = Document(str(template_path))
    except Exception as e:
        logger.exception("read docx first table preview failed")
        return {"success": False, "message": f"无法读取 Word：{e}"}
    if not doc.tables:
        return {"success": False, "message": "文档中无表格"}
    table = doc.tables[0]
    if len(table.rows) < 1:
        return {"success": False, "message": "表格为空"}
    headers_raw = _table_row_texts(table.rows[0])
    headers: list[str] = []
    for i, h in enumerate(headers_raw):
        name = (h or "").strip()
        if not name:
            name = f"列{i + 1}"
        headers.append(name)
    sample_row: dict[str, Any] = {}
    if len(table.rows) >= 2:
        vals = _table_row_texts(table.rows[1])
        for i, h in enumerate(headers):
            sample_row[h] = vals[i] if i < len(vals) else ""
    else:
        for h in headers:
            sample_row[h] = ""
    return {
        "success": True,
        "headers": headers,
        "sample_rows": [sample_row],
        "sheet_name": sheet_name,
    }


def read_price_list_docx_table_preview(template_path: Path) -> dict[str, Any]:
    """
    读取报价模板 ``Document`` 中**首个表格**的第 1 行（表头）与第 2 行（示例数据，若存在），
    与 ``build_price_list_docx_bytes`` 写入的首表一致，供模板预览页与「导出价格表」对齐。
    """
    return read_docx_first_table_preview(template_path, sheet_name="Word 报价表（首表）")


def price_list_docx_template_public_hint(template_path: Path) -> str:
    """用于 UI 展示的相对仓库路径；无法相对化时退回文件名。"""
    try:
        root = _repo_root()
        resolved = template_path.resolve()
        root_res = root.resolve()
        rel = resolved.relative_to(root_res)
        return str(rel).replace("\\", "/")
    except Exception:
        return template_path.name


def build_price_list_template_preview_json(template_slug: str | None = None) -> dict[str, Any]:
    tpl = resolve_price_list_docx_template(template_slug)
    if not tpl:
        return {
            "success": False,
            "message": (
                "未找到 Word 模板：请将默认模板放在 424/document_templates/price_list_default.docx，"
                "或通过管理接口登记；亦可设置环境变量 FHD_PRICE_LIST_DOCX_TEMPLATE。"
            ),
        }
    data = read_price_list_docx_table_preview(tpl)
    if not data.get("success"):
        return data
    data["template_hint"] = price_list_docx_template_public_hint(tpl)
    return data


def build_sales_contract_template_preview_json(template_slug: str | None = None) -> dict[str, Any]:
    """
    与 ``POST /api/sales-contract/generate`` 同源解析模板；Word 抽首表，Excel 抽首工作表表头与示例行。
    """
    from backend.document_template_service import ROLE_SALES_CONTRACT, resolve_template_path_with_meta

    slug = (template_slug or "").strip() or None
    tpl_path, _eff = resolve_template_path_with_meta(role=ROLE_SALES_CONTRACT, slug=slug)
    if not tpl_path:
        return {
            "success": False,
            "message": (
                "未找到销售合同模板：请在 PostgreSQL document_templates 登记 role=sales_contract_docx（.xls/.xlsx/.docx），"
                "或设置环境变量 FHD_SALES_CONTRACT_TEMPLATE。"
            ),
        }
    suf = tpl_path.suffix.lower()
    if suf in (".xls", ".xlsx", ".xlsm"):
        from backend.sales_contract_excel_generate import read_excel_sales_contract_preview

        data = read_excel_sales_contract_preview(tpl_path)
    else:
        data = read_docx_first_table_preview(tpl_path, sheet_name="Word 销售合同（首表）")
    if not data.get("success"):
        return data
    data["template_hint"] = price_list_docx_template_public_hint(tpl_path)
    return data


def build_price_list_docx_bytes(
    template_path: Path,
    *,
    customer_name: str,
    quote_date: str | None,
    products: list[dict[str, Any]],
) -> bytes:
    qd = (quote_date or "").strip() or date.today().strftime("%Y-%m-%d")
    doc = Document(str(template_path))
    _fill_customer_and_quote_date(doc, customer_name, qd)
    _fill_first_table_products(doc, products, start_row=1)
    if doc.tables:
        _center_table_cell_text(doc.tables[0])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
