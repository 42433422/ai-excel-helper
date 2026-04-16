from docx import Document
import os
from datetime import datetime

def create_sales_contract(
    template_path: str,
    output_path: str,
    customer_name: str = "深圳市百木鼎家具有限公司",
    customer_phone: str = "",
    contract_date: str = None,
    return_buckets_expected: int = 0,
    return_buckets_actual: int = 0,
    products: list = None
):
    """
    创建销售合同及送货凭证
    """
    doc = Document(template_path)

    if contract_date is None:
        now = datetime.now()
        contract_date = f"{now.year}年{now.month:02d}月{now.day:02d}日"

    for para in doc.paragraphs:
        full_text = para.text
        if "客户：" in full_text and "MESSRS" not in full_text:
            for run in para.runs:
                if "客户：" in run.text:
                    if "惠州市宝盈家具有限公司" in run.text:
                        run.text = f"客户：{customer_name}"
                    else:
                        run.text = run.text.replace("客户：    ", f"客户：{customer_name}")
                    if customer_phone:
                        run.text += f" 电话：{customer_phone}"

        if "MESSRS" in full_text and "TEL" in full_text:
            for run in para.runs:
                if "MESSRS" in run.text:
                    run.text = f"{customer_name}\t"
                if "TEL" in run.text:
                    run.text = f"电话：{customer_phone}\t" if customer_phone else "电话：\t"

        if full_text.strip().startswith("ADDRESS") or ("ADDRESS" in full_text and "DATE" in full_text):
            for run in para.runs:
                if "ADDRESS" in run.text:
                    run.text = run.text.replace("ADDRESS", customer_name)
                if "DATE" in run.text:
                    run.text = run.text.replace("DATE", contract_date)

    table = doc.tables[0]

    if products is None:
        products = [
            {
                "model_number": "306B",
                "name": "PU亮光硬化剂",
                "spec": "10KG×1",
                "unit": "桶",
                "quantity": "10 KG",
                "unit_price": "39.2",
                "amount": "392"
            }
        ]

    for idx, product in enumerate(products):
        if idx + 1 < len(table.rows):
            row = table.rows[idx + 1]

            if product.get("model_number"):
                row.cells[0].text = product["model_number"]

            if product.get("name"):
                model_no = product.get("model_number", "")
                product_name = product["name"]
                row.cells[1].text = f"{model_no}\n{product_name}" if model_no else product_name

            if product.get("spec"):
                row.cells[2].text = product["spec"]

            if product.get("unit"):
                row.cells[3].text = product["unit"]
                row.cells[4].text = product["unit"]

            if product.get("quantity"):
                row.cells[5].text = product["quantity"]
                row.cells[6].text = product["quantity"]

            if product.get("unit_price"):
                row.cells[7].text = product["unit_price"]
                row.cells[8].text = "元/KG"

            if product.get("amount"):
                row.cells[9].text = f"{product['amount']}元"

    total_quantity = 0.0
    total_amount = 0.0
    for p in products:
        if p.get("quantity"):
            try:
                qty = float(p.get("quantity", "0").replace(" KG", ""))
                total_quantity += qty
            except:
                pass
        if p.get("amount"):
            try:
                total_amount += float(p.get("amount", "0"))
            except:
                pass

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if "3 KG" in cell.text:
                cell.text = cell.text.replace("3 KG", f"{total_quantity} KG")
            if "应退桶(  )" in cell.text:
                cell.text = cell.text.replace("(  )", f"({return_buckets_expected})")
            if "实退桶(  )" in cell.text:
                cell.text = cell.text.replace("(  )", f"({return_buckets_actual})")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    print(f"销售合同已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    template = r"e:\FHD\424\转 Word_扫描全能王 11.04.26 13.31.docx"
    output = r"e:\FHD\深圳市百木鼎家具有限公司_306B销售合同.docx"

    create_sales_contract(
        template_path=template,
        output_path=output,
        customer_name="深圳市百木鼎家具有限公司",
        customer_phone="",
        contract_date="2026年04月11日",
        return_buckets_expected=1,
        return_buckets_actual=0,
        products=[
            {
                "model_number": "306B",
                "name": "PU亮光硬化剂",
                "spec": "10KG×1",
                "unit": "桶",
                "quantity": "10 KG",
                "unit_price": "39.2",
                "amount": "392"
            }
        ]
    )